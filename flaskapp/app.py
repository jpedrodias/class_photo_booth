# Standard library imports
import base64
import csv
import os
import re
import shutil
import traceback
import zipfile
import time
import pickle, json
from datetime import datetime, timedelta, timezone
from functools import wraps
from io import BytesIO, StringIO
from ipaddress import ip_address

# Third party imports
import redis
import msgpack

# Redis Queue imports
from rq import Queue
from rq.job import Job
#from rq.job import Job

# OpenCV and numpy for image processing
import cv2
import numpy as np

# DOCX imports
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# Flask imports
from flask import Flask, jsonify, render_template, request, redirect, url_for, send_file, session, Response, make_response, flash
from flask_mail import Mail, Message
from flask_session import Session
from flask_sqlalchemy import SQLAlchemy
from flask_wtf.csrf import CSRFProtect
from werkzeug.security import generate_password_hash, check_password_hash, safe_join
from werkzeug.utils import secure_filename
from itsdangerous import URLSafeTimedSerializer



app = Flask(__name__)
app.config.from_object('config.DevelopmentConfig')  # Use DevelopmentConfig by default  

# Inicializar extensão Flask-Session
Session(app)

# Inicializar serializer para tokens
serializer = URLSafeTimedSerializer(app.secret_key)


# Inicializar proteção CSRF
csrf = CSRFProtect(app)

# Disponibilizar token CSRF globalmente nos templates
@app.context_processor
def inject_csrf_token():
    from flask_wtf.csrf import generate_csrf
    return dict(csrf_token=generate_csrf)
# End inject_csrf_token


# Inicializar SQLAlchemy
db = SQLAlchemy(app)

# Configurar Flask-Mail
mail = Mail(app)

# Configurar Redis Queue
redis_conn = redis.from_url(app.config['RQ_REDIS_URL'])
email_queue = Queue('email', connection=redis_conn, default_timeout=app.config['RQ_DEFAULT_TIMEOUT'])

# Configuração das pastas
BASE_DIR = app.config['BASE_DIR'] 
PHOTOS_DIR = app.config['PHOTOS_DIR']
THUMBS_DIR = app.config['THUMBS_DIR']
DEBUG = app.config['DEBUG']

#DEBUG Environment
print("DEBUG mode is set to:", DEBUG)
print("Base directory is set to:", BASE_DIR)
print('Email settings:')
print(f"MAIL_USERNAME: {app.config['MAIL_USERNAME']}")


# ========== Modelos SQLAlchemy ==========

class AddUserSecurityCheck():
    # to be inherited by PreUser and User
    pattern_email = re.compile(r'([A-Za-z0-9]+[.-_])*[A-Za-z0-9]+@[A-Za-z0-9-]+(\.[A-Z|a-z]{2,})+')
    pattern_password = "^(?=.*?[A-Z])(?=.*?[a-z])(?=.*?[0-9])(?=.*?[#?!@$%^&*-]).{8,}$"

    @classmethod
    def validate_email_format(cls, email_to_test: str) -> bool:
        """Validate email format using regex - can be called without instantiation"""
        return bool(re.fullmatch(cls.pattern_email, email_to_test))

    @classmethod
    def validate_password_strength(cls, password: str) -> tuple[bool, str]:
        """Validate password strength - at least 6 chars with mixed case and numbers"""
        if len(password) < 6:
            return False, "Password deve ter pelo menos 6 caracteres."
        
        if not re.search(r'[A-Z]', password):
            return False, "Password deve conter pelo menos uma letra maiúscula."
        
        if not re.search(r'[a-z]', password):
            return False, "Password deve conter pelo menos uma letra minúscula."
        
        if not re.search(r'\d', password):
            return False, "Password deve conter pelo menos um número."
        
        return True, "Password válida."

    def is_valid_email(self, email_to_test: str) -> bool:
        """Instance method that calls the class method for backward compatibility"""
        return self.validate_email_format(email_to_test)

    def is_valid_password(self, password_to_test: str) -> bool:
        if re.match(AddUserSecurityCheck.pattern_password, password_to_test):
            return True
        else:
            return False
#End class AddUserSecurityCheck


class PreUser(db.Model, AddUserSecurityCheck):
    __tablename__ = 'pre_users' 
    
    id = db.Column(db.Integer, primary_key=True)
    # Remember: to QUERY use "_email" and not "email"; to Instantiate use "email" and not "_email"
    _email = db.Column('email', db.String(120), unique=True, nullable=False)
    code = db.Column(db.String(120), unique=False, nullable=False)
    date = db.Column(db.DateTime, unique=False, nullable=False, default=lambda: datetime.now(timezone.utc))
    reason = db.Column(db.String(255), unique=False, nullable=True, default='')  # Motivo opcional para o pré-registo
    email_job_id = db.Column(db.String(36), unique=False, nullable=True)  # ID da tarefa de envio de email

    def __repr__(self):
        return f'<PreUser {self.email}>'

    @property
    def email(self):
        return self._email
    
    @email.getter
    def email(self):
        return self._email

    @email.setter
    def email(self, value: str) -> None:
        if self.is_valid_email(value):
            self._email = value
        else:
            raise ValueError('Invalid email')
    
    @staticmethod
    def create_random_code(length: int = 6) -> str:
        """Returns random secret numbers with letters and numbers"""
        from string import ascii_lowercase, digits
        from random import sample
        return ''.join(sample(ascii_lowercase + digits, length))
#end class PreUser


class User(db.Model, AddUserSecurityCheck):
    __tablename__ = 'users'
    
    id = db.Column(db.Integer, primary_key=True)
    _email = db.Column('username', db.String(120), unique=True, nullable=False)
    password_hash = db.Column('password', db.String(180), unique=False, nullable=False)
    name = db.Column('name', db.String(120), unique=False, nullable=False)
    role = db.Column(db.String(20), nullable=False, default='none')  # none, viewer, editor, admin
    is_verified = db.Column(db.Boolean, nullable=False, default=True)  # Por defeito verificado
    logins = db.relationship('LoginLog', backref=db.backref('user', lazy='joined'))

    def __repr__(self):
        return f'<User {self.email}>'
    
    @property
    def email(self):
        return self._email
    
    @email.getter
    def email(self):
        return self._email

    @email.setter
    def email(self, value):
        if self.is_valid_email(value):
            self._email = value
        else:
            raise ValueError('Invalid email')
    
    @property
    def password(self):
        return self.password_hash
    
    @password.setter
    def password(self, value):
        # Para registo, aceitar passwords mais simples (mínimo 6 caracteres)
        if len(value) >= 6:
            self.password_hash = generate_password_hash(value)
        else:
            raise ValueError('Password deve ter pelo menos 6 caracteres.')
    
    def check_password(self, value):
        return check_password_hash(self.password_hash, value)
    
    def has_permission(self, permission):
        """Check if user has specific permission"""
        permissions = {
            'none': [],
            'viewer': ['view_turmas', 'view_photos'],
            'editor': ['view_turmas', 'view_photos', 'capture_photos', 'manage_students'],
            'admin': ['view_turmas', 'view_photos', 'capture_photos', 'manage_students', 'edit_classes', 'manage_turmas', 'upload_csv', 'system_nuke']
        }
        return permission in permissions.get(self.role, [])
#end class User


class Turma(db.Model):
    __tablename__ = 'turmas'
    
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(100), nullable=False)  # Nome original (display)
    nome_seguro = db.Column(db.String(100), nullable=False, unique=True)  # Nome sanitizado (filesystem)
    nome_professor = db.Column(db.String(100), nullable=False, unique=False, default='')  # Nome do professor responsável
    email_professor = db.Column(db.String(255), nullable=False, default='')  # Email do professor responsável
    last_updated = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))  # Data da última atualização
    alunos = db.relationship('Aluno', backref='turma', lazy=True, cascade='all, delete-orphan')
    
    def __init__(self, nome, **kwargs):
        """Construtor que automaticamente gera nome seguro"""
        super(Turma, self).__init__(**kwargs)
        self.nome = nome
        self.nome_seguro = self._generate_safe_name(nome)
    
    def _generate_safe_name(self, nome):
        """Gera nome seguro usando secure_filename do Flask"""
        nome_seguro = secure_filename(nome)
        
        # Garantir que não fica vazio
        if not nome_seguro or nome_seguro.isspace():
            nome_seguro = f"turma_{self.id or 'nova'}"
        
        # Limitar comprimento
        nome_seguro = nome_seguro[:50]
        
        # Garantir unicidade (apenas se já existe uma sessão ativa)
        if db.session:
            counter = 1
            original_nome = nome_seguro
            # Excluir a própria turma da verificação de unicidade
            query = Turma.query.filter_by(nome_seguro=nome_seguro)
            if self.id:  # Se a turma já tem ID, excluir ela própria
                query = query.filter(Turma.id != self.id)
            
            while query.first():
                nome_seguro = f"{original_nome}_{counter}"
                counter += 1
                # Atualizar a query para o novo nome
                query = Turma.query.filter_by(nome_seguro=nome_seguro)
                if self.id:
                    query = query.filter(Turma.id != self.id)
        return nome_seguro
    
    def get_safe_directory_path(self, base_dir):
        """Retorna caminho seguro para diretório da turma"""
        try:
            return safe_join(base_dir, self.nome_seguro)
        except ValueError:
            # Se mesmo assim for perigoso, usar ID
            safe_name = f"turma_{self.id}"
            return safe_join(base_dir, safe_name)
    
    def get_foto_directory(self):
        """Retorna diretório das fotos originais"""
        return self.get_safe_directory_path(PHOTOS_DIR)
    
    def get_thumb_directory(self):
        """Retorna diretório das thumbnails"""
        return self.get_safe_directory_path(THUMBS_DIR)
    
    def create_directories(self):
        """Cria diretórios da turma se não existirem"""
        foto_dir = self.get_foto_directory()
        thumb_dir = self.get_thumb_directory()
        
        safe_makedirs(foto_dir)
        safe_makedirs(thumb_dir)
        
        return foto_dir, thumb_dir
    
    def update_nome(self, novo_nome):
        """Atualiza nome da turma e renomeia diretórios se necessário"""
        # Se o nome não mudou, não fazer nada
        if self.nome == novo_nome:
            return
        
        nome_antigo_seguro = self.nome_seguro
        novo_nome_seguro = self._generate_safe_name(novo_nome)
        
        # Se o nome seguro mudou, renomear diretórios
        if nome_antigo_seguro != novo_nome_seguro:
            self._rename_directories(nome_antigo_seguro, novo_nome_seguro)
        
        self.nome = novo_nome
        self.nome_seguro = novo_nome_seguro
    
    def _rename_directories(self, old_safe_name, new_safe_name):
        """Renomeia diretórios quando nome seguro muda"""
        import shutil
        
        try:
            # Renomear diretório de fotos
            old_foto_dir = safe_join(PHOTOS_DIR, old_safe_name)
            new_foto_dir = safe_join(PHOTOS_DIR, new_safe_name)
            if os.path.exists(old_foto_dir):
                shutil.move(old_foto_dir, new_foto_dir)
            
            # Renomear diretório de thumbnails
            old_thumb_dir = safe_join(THUMBS_DIR, old_safe_name)
            new_thumb_dir = safe_join(THUMBS_DIR, new_safe_name)
            if os.path.exists(old_thumb_dir):
                shutil.move(old_thumb_dir, new_thumb_dir)
        except (ValueError, OSError) as e:
            print(f"Erro ao renomear diretórios: {e}")
    
    def delete_directories(self):
        """Remove diretórios da turma"""
        import shutil
        
        try:
            foto_dir = self.get_foto_directory()
            thumb_dir = self.get_thumb_directory()
            
            if os.path.exists(foto_dir):
                shutil.rmtree(foto_dir)
            
            if os.path.exists(thumb_dir):
                shutil.rmtree(thumb_dir)
        except (ValueError, OSError) as e:
            print(f"Erro ao remover diretórios: {e}")
    
    @classmethod
    def create_safe(cls, nome):
        """Factory method que cria turma com nome seguro"""
        turma = cls(nome=nome)
        db.session.add(turma)
        db.session.flush()  # Para obter ID antes do commit
        
        # Recriar nome seguro com ID se necessário
        if not turma.nome_seguro or turma.nome_seguro == "turma_nova":
            turma.nome_seguro = turma._generate_safe_name(nome)
        
        turma.create_directories()
        return turma
    
    @classmethod
    def get_nome_seguro_by_nome(cls, nome):
        """Retorna nome_seguro de uma turma pelo seu nome"""
        turma = cls.query.filter_by(nome=nome).first()
        return turma.nome_seguro if turma else nome
    
    def update_last_modified(self):
        """Atualiza o timestamp de última modificação"""
        self.last_updated = datetime.now(timezone.utc)
        
    def __repr__(self):
        return f'<Turma {self.nome} ({self.nome_seguro})>'
# End class Turma


class Aluno(db.Model):
    __tablename__ = 'alunos'
    
    id = db.Column(db.Integer, primary_key=True)
    processo = db.Column(db.String(50), nullable=False, unique=True, index=True)  # Único globalmente
    nome = db.Column(db.String(200), nullable=False)
    numero = db.Column(db.Integer, nullable=True)  # Número na turma (pode repetir)
    email = db.Column(db.String(255), nullable=False, default='')  # Email do aluno
    autorizacao = db.Column(db.Boolean, default=True, nullable=False)  # Autorização para redes sociais
    notes = db.Column(db.Text, nullable=True)  # Notas sobre o aluno
    foto_existe = db.Column(db.Boolean, default=False, nullable=False)  # Nova propriedade
    foto_tirada = db.Column(db.Boolean, default=False, nullable=False)
    turma_id = db.Column(db.Integer, db.ForeignKey('turmas.id'), nullable=False)
    
    @classmethod
    def validate_and_sanitize_processo(cls, processo):
        """
        Valida e sanitiza o número do processo.
        Retorna (processo_validado, erro_msg) onde erro_msg é None se válido.
        """
        if not processo or not processo.strip():
            return None, "Número do processo é obrigatório!"
        
        processo = processo.strip()
        
        # Verificar se é um número inteiro
        try:
            processo_int = int(processo)
            if processo_int <= 0:
                return None, "O número do processo deve ser um número inteiro positivo!"
            
            # Converter de volta para string para manter consistência
            return str(processo_int), None
            
        except ValueError:
            return None, "O número do processo deve ser um número inteiro válido (ex: NIF, número de estudante, etc.)!"
    
    def __repr__(self):
        return f'<Aluno {self.nome} ({self.processo})>'
# End class Aluno


class LoginLog(db.Model):
    __tablename__ = 'logs_login'
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.DateTime, nullable=False,  default=lambda: datetime.now(timezone.utc))
    success = db.Column(db.Boolean, unique=False, nullable=False)
    remote_addr = db.Column('ip_address', db.String(64), unique=False, nullable=False) # request.remote_addr
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))

    @classmethod
    def check_failed_logins(cls, ip_address, max_attempts=5):
        """Verifica tentativas falhadas de login e bane IP se necessário"""
        
        # Verificar tentativas nas últimas 15 minutos
        time_threshold = datetime.now(timezone.utc) - timedelta(minutes=1)
        failed_attempts = cls.query.filter(
            cls.remote_addr == ip_address,
            cls.success == False,
            cls.date >= time_threshold
        ).count()
        
        if failed_attempts >= max_attempts:
            BannedIPs.ban_ip(ip_address)
            return True
        return False

    @classmethod
    def log_attempt(cls, user_id, ip_address, success):
        """Regista tentativa de login"""
        login_log = cls(
            user_id=user_id,
            remote_addr=ip_address,
            success=success
        )
        db.session.add(login_log)
        db.session.commit()
#End class LoginLog


class BannedIPs(db.Model):
    """Anti Brute Force Attacks - based on the ip address"""
    __tablename__ = 'banned_ips'
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.DateTime, nullable=False, default=lambda: datetime.now(timezone.utc))

    #define INET4_ADDRSTRLEN 16 + ...
    #define INET6_ADDRSTRLEN 46 + date
    remote_addr = db.Column('ip_address', db.String(64), unique=False, nullable=False, index=True) # request.remote_addr

    @staticmethod
    def normalize_ip(raw_ip: str):
        if not raw_ip:
            return None
        s = raw_ip.strip()
        # alguns proxies/URLs passam IPv6 entre [ ... ]
        if s.startswith('[') and s.endswith(']'):
            s = s[1:-1]
        ip = ip_address(s)
        # converter IPv4-mapped (::ffff:1.2.3.4) para IPv4 “puro”
        if ip.version == 6 and getattr(ip, "ipv4_mapped", None):
            ip = ip.ipv4_mapped
        return ip.compressed  # forma canónica (zeros comprimidos, minúsculas)

    @classmethod
    def is_banned(cls, ip_address):
        """Verifica se um IP está banido"""
        ip_norm = cls.normalize_ip(ip_address)
        if ip_norm is None :
            return False
        #banned_ip = cls.query.filter_by(remote_addr=ip_norm).first()
        #return banned_ip is not None
        return db.session.query(cls.id).filter_by(remote_addr=ip_norm).first() is not None
    
    @classmethod
    def ban_ip(cls, ip_address):
        """Bane um IP por tentativas excessivas de login"""
        ip_norm = cls.normalize_ip(ip_address)
        if ip_norm is None:
            return False
        if not cls.is_banned(ip_norm):
            db.session.add(cls(remote_addr=ip_norm))
            db.session.commit()
            return True
        return False
# End class BannedIPs

# ========== TOOLS ==========

# returns a String, example 2021 ou 20/21
def calc_School_Year(day=None, formato='{}{}', formato_size=2):
    valid_input = False
    if isinstance(day, type(None)):
        day = datetime.today()
    
    if isinstance(day, str):
        try:
            day = datetime.strptime(day, '%Y-%m-%d')
        except:
            day = None
            
    if not isinstance(day, datetime):
        raise TypeError('Expected datetime object or string "%Y-%m-%d" ')
    
    m = day.month
    y = day.year
    if m >= 8:
        y1, y2 = y, y +1
    else:
        y1, y2 = y - 1, y
    
    y1 = f'{y1}'[formato_size:]
    y2 = f'{y2}'[formato_size:]
    return formato.format(y1, y2)
#End def calc_School_Year


# Funções auxiliares para autenticação e segurança
def get_current_user():
    """Obtém o utilizador atual da sessão e renova TTL se sliding_expiry=True"""
    user_id = session.get('user_id')
    if not user_id:
        return None
    
    # Verificar expiração
    exp_ts = session.get('expires_at')
    if exp_ts:
        try:
            if datetime.now(timezone.utc).timestamp() > float(exp_ts):
                session.clear()
                return None
        except Exception:
            session.clear()
            return None
    
    # Atualizar last_seen apenas para requests não-API (evitar atualizar constantemente)
    # Não atualizar last_seen para requests de API/status para não interferir com monitorização
    if not request.path.startswith('/settings/redis/'):
        session['last_seen'] = datetime.now(timezone.utc).isoformat()
    
    # Se sliding_expiry=True, renovar TTL usando configuração
    if session.get('sliding_expiry', False):
        new_expiry = (datetime.now(timezone.utc) + app.config['TEMPORARY_SESSION_LIFETIME']).timestamp()
        session['expires_at'] = new_expiry
    
    return db.session.get(User, user_id)
# End def get_current_user


# Decorator para rotas que requerem login
def required_login(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        user = get_current_user()
        if not user:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function
# End def required_login


# Decorator para verificar role mínimo
def required_role(min_role):
    """
    Decorator que verifica se o utilizador tem pelo menos o role especificado.
    Hierarquia: none < viewer < editor < admin
    """
    role_hierarchy = {
        'none':   0,
        'viewer': 1, 
        'editor': 2,
        'admin':  3
    }
    
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            user = get_current_user()
            if not user:
                return redirect(url_for('login'))
            
            user_level = role_hierarchy.get(user.role, 0)
            required_level = role_hierarchy.get(min_role, 0)
            
            if user_level < required_level:
                if user.role == 'none':
                    flash('A sua conta está à espera de validação pelo administrador. Não tem permissões para aceder a esta funcionalidade.', 'info')
                else:
                    flash(f'Precisa de pelo menos permissões de {min_role} para aceder a esta funcionalidade.', 'error')
                return redirect(url_for('index'))
            
            return f(*args, **kwargs)
        return decorated_function
    return decorator
# End def required_role


# Certifique-se de que as pastas existem com permissões adequadas
def create_directories_with_permissions():
    for directory in [PHOTOS_DIR, THUMBS_DIR]:
        if not safe_makedirs(directory, verbose=True):
            print(f"Erro crítico: Não foi possível criar diretório {directory}")
# End def create_directories_with_permissions


def safe_makedirs(directory, verbose=False):
    """Cria diretórios de forma segura, compatível com Windows e Linux"""
    try:
        os.makedirs(directory, exist_ok=True)
        # Apenas tenta definir permissões em sistemas Unix/Linux
        if os.name != 'nt':
            try:
                os.chmod(directory, 0o777)
            except (PermissionError, OSError) as e:
                if verbose:
                    print(f"Aviso: Não foi possível definir permissões para {directory}: {e}")
        return True
    except PermissionError as e:
        if verbose:
            print(f"Erro de permissão ao criar diretório {directory}: {e}")
        return False
    except Exception as e:
        if verbose:
            print(f"Erro ao criar diretório {directory}: {e}")
        return False
# End def safe_makedirs


# Funções para manipulação de documentos DOCX
def docx_replace(doc, dicionario):
    """Substitui placeholders no documento DOCX preservando a formatação"""
    # Substituir em parágrafos
    for paragraph in doc.paragraphs:
        for key, value in dicionario.items():
            placeholder = f'{{{key}}}'
            if placeholder in paragraph.text:
                # Preservar formatação dos runs
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in dicionario.items():
                        placeholder = f'{{{key}}}'
                        if placeholder in paragraph.text:
                            # Preservar formatação dos runs
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    
    # Substituir em headers e footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for key, value in dicionario.items():
                    placeholder = f'{{{key}}}'
                    if placeholder in paragraph.text:
                        # Preservar formatação dos runs
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for key, value in dicionario.items():
                    placeholder = f'{{{key}}}'
                    if placeholder in paragraph.text:
                        # Preservar formatação dos runs
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
# End docx_replace


def process_image_for_docx(image_path, target_width_cm, target_height_cm):
    """Processa imagem para o DOCX: redimensiona mantendo proporção e faz crop"""
    try:
        import cv2
        from PIL import Image
        from io import BytesIO
        
        # Ler imagem
        if image_path.endswith('.jpg') or image_path.endswith('.jpeg'):
            # Usar OpenCV para ler
            img = cv2.imread(image_path)
            if img is None:
                return None
            # Converter de BGR para RGB
            img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            # Converter para PIL
            img_pil = Image.fromarray(img)
        else:
            # Usar PIL diretamente
            img_pil = Image.open(image_path)
        
        # Converter dimensões de cm para pixels (assumindo 150 DPI)
        dpi = 150
        target_width_px = int(target_width_cm * dpi / 2.54)
        target_height_px = int(target_height_cm * dpi / 2.54)
        
        # Obter dimensões originais
        original_width, original_height = img_pil.size
        
        # Calcular proporções
        width_ratio = target_width_px / original_width
        height_ratio = target_height_px / original_height
        
        # Usar a maior proporção para garantir que a imagem cubra completamente o target
        scale_ratio = max(width_ratio, height_ratio)
        
        # Redimensionar
        new_width = int(original_width * scale_ratio)
        new_height = int(original_height * scale_ratio)
        img_resized = img_pil.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Calcular posição para crop central
        left = (new_width - target_width_px) // 2
        top = (new_height - target_height_px) // 2
        right = left + target_width_px
        bottom = top + target_height_px
        
        # Fazer crop
        img_cropped = img_resized.crop((left, top, right, bottom))
        
        # Salvar em BytesIO
        output = BytesIO()
        img_cropped.save(output, format='JPEG', quality=85)
        output.seek(0)
        
        return output
        
    except Exception as e:
        print(f"Erro ao processar imagem {image_path}: {e}")
        return None
# End def process_image_for_docx


def create_docx_with_photos(turma_nome):
    """Cria documento DOCX com fotos dos alunos da turma"""
    
    DEFAULT_FONT = 'Calibri Light' # Definindo a fonte padrão como uma constante
    # Buscar turma na base de dados
    turma_obj = Turma.query.filter_by(nome=turma_nome).first()
    if not turma_obj:
        return None
    
    # Caminho do template
    template_path = os.path.join(BASE_DIR, 'templates', 'template_relacao_alunos_fotos.docx')
    if not os.path.exists(template_path):
        return None
    
    # Criar documento a partir do template
    doc = Document(template_path)
    
    # Definir propriedades do documento
    doc.core_properties.author = 'Class Photo Booth by Pedro Dias'
    doc.core_properties.title = f'Relação de Alunos - {turma_nome}'
    doc.core_properties.subject = f'Relação de Alunos - {turma_nome}'
    
    # Dicionário para substituições
    current_date = datetime.now().strftime('%d/%m/%Y')
    
    dicionario = {
        'turma': turma_nome,
        'ano_letivo': calc_School_Year(turma_obj.last_updated.strftime('%Y-%m-%d'), '{}/{}'),
        'date': turma_obj.last_updated.strftime('%d/%m/%Y %Hh%M'),
        'professor': turma_obj.nome_professor  # Pode ser personalizado
    }
    
    # Substituir placeholders no template
    docx_replace(doc, dicionario)
    
    # Obter alunos da turma ordenados
    alunos_ordenados = sorted(
        turma_obj.alunos, 
        key=lambda x: (x.numero is None, x.numero if x.numero is not None else 0, x.nome)
    )
    
    # Incluir todos os alunos (com e sem fotos)
    if not alunos_ordenados:
        return None
    
    # Configurar tabela
    NUM_PER_ROW = 4
    table = doc.add_table(rows=0, cols=NUM_PER_ROW)
    table.style = 'Table Grid'
    table.width = Cm(20)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configurar largura das colunas
    for i in range(NUM_PER_ROW):
        table.columns[i].alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[i].width = Cm(20/NUM_PER_ROW)
    
    # Determinar tamanho das imagens baseado no número total de alunos
    num_alunos = len(alunos_ordenados)
    if num_alunos > 28:
        img_width, img_height = 3.47, 2.6 # 3.6, 2.7
    else:
        img_width, img_height = 4, 3
    
    # Adicionar fotos dos alunos à tabela
    for index, aluno in enumerate(alunos_ordenados):
        # Adicionar nova linha a cada NUM_PER_ROW alunos
        if index % NUM_PER_ROW == 0:
            row = table.add_row().cells
        
        coluna = index % NUM_PER_ROW
        
        # Caminho da thumbnail (imagem otimizada)
        thumb_path = os.path.join(turma_obj.get_thumb_directory(), f'{aluno.processo}.jpg')
        placeholder_path = os.path.join(BASE_DIR, 'static', 'student_icon.jpg')
        
        # Usar foto do aluno se existir, senão usar placeholder
        image_path = thumb_path if os.path.exists(thumb_path) else placeholder_path
        
        if os.path.exists(image_path):
            # Adicionar imagem à célula
            paragraph = row[coluna].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            
            try:
                # Processar imagem para ajustar dimensões e fazer crop
                processed_image = process_image_for_docx(image_path, img_width, img_height)
                
                if processed_image:
                    run.add_picture(processed_image, width=Cm(img_width), height=Cm(img_height))
                else:
                    # Fallback: usar imagem original
                    run.add_picture(image_path, width=Cm(img_width), height=Cm(img_height))
                
                run.add_break()
                
                # Adicionar texto com número e nome formatado
                numero_display = aluno.numero if aluno.numero else index + 1
                text = f'{numero_display} {aluno.nome}'
                text_size = 12 if len(text) < 24 else 10
                text_run = paragraph.add_run(text)
                # Aplicar formatação Times New Roman 16pt
                text_run.font.name = DEFAULT_FONT 
                text_run.font.size = Pt(text_size)
            except Exception as e:
                # Se houver erro ao adicionar imagem, adicionar apenas o texto
                numero_display = aluno.numero if aluno.numero else index + 1
                text = f'{numero_display} {aluno.nome}'
                text_size = 12 if len(text) < 24 else 10
                text_run = paragraph.add_run(text)
                # Aplicar formatação Times New Roman 16pt
                text_run.font.name = DEFAULT_FONT 
                text_run.font.size = Pt(text_size)
                print(f"Erro ao adicionar imagem para {aluno.nome}: {e}")
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Salvar documento na memória
    memory_file = BytesIO()
    doc.save(memory_file)
    memory_file.seek(0)
    
    return memory_file
        
    #except Exception as e:
    #    print(f"Erro ao criar documento DOCX: {e}")
    #    return None
# End def create_docx_with_photos


# ========== Flask Routes ==========
def no_cache(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        response = f(*args, **kwargs)
        if isinstance(response, str):
            response = make_response(response)
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response
    return wrapper
# End def no_cache (decorator)


@app.route('/')
@required_login
def index():
    user = get_current_user()
    
    # Se o utilizador não tem permissões, mostrar página de boas-vindas
    if user.role == 'none':
        return render_template('home.html', user=user, current_user=user)
    
    # Para utilizadores com permissões, redirecionar para turmas
    return redirect(url_for('turmas'))
# End def index (homepage)


@app.route('/login/', methods=['GET', 'POST'])
@app.route('/login/<action_url>', methods=['GET', 'POST'])
def login(action_url=None):
    # Verificar se o IP está banido
    ip_address = request.headers.get('Cf-Connecting-Ip', None) # NEED if client/server is behind Cloudflare Proxy
    if not ip_address:
        ip_address = request.access_route[-1] # request.remote_addr

    # FUNCIONALIDADE DESATIVADA: Banning IPs for excessive login attempts
    DESATIVAR = False
    if BannedIPs.is_banned(ip_address) and not DESATIVAR:
        flash('O seu IP foi bloqueado devido a tentativas excessivas de login. Contacte o administrador.', 'error')
        return render_template('login.html', action='login')
    
    if request.method == 'POST':
        action = action_url or 'login'
        
        if action == 'login':
            email = request.form.get('email', '').strip()
            password = request.form.get('password', '')
            remember_me = request.form.get('remember_me') == 'on'
            
            if not email or not password:
                flash('Email e palavra-passe são obrigatórios.', 'error')
                return render_template('login.html', action='login')
            
            # Validar formato do email
            if not AddUserSecurityCheck.validate_email_format(email):
                flash('Formato de email inválido.', 'error')
                return render_template('login.html', action='login')
            
            # Verificar utilizador
            user = User.query.filter_by(_email=email).first()
            
            if not user:
                # Log da tentativa mesmo sem utilizador para detectar ataques
                LoginLog.log_attempt(None, ip_address, False)
                
                # Verificar se deve banir o IP
                if LoginLog.check_failed_logins(ip_address):
                    flash('Demasiadas tentativas falhadas. O seu IP foi bloqueado.', 'error')
                else:
                    flash('Email ou palavra-passe incorretos.', 'error')
                
                return render_template('login.html', action='login')
            
            # Verificar se a conta está verificada
            if not user.is_verified:
                flash('Conta não verificada. Verifique o seu email ou registe-se novamente.', 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Verificar se o utilizador tem password definida
            if not user.password_hash:
                flash('Password não definida. Defina a sua password primeiro.', 'error')
                return render_template('login.html', action='login', email=email)
            
            if user.check_password(password):
                # Login bem-sucedido
                session['user_id'] = user.id
                session['user_email'] = user.email
                session['user_role'] = user.role
                
                # Configurar duração da sessão baseado no "remember me"
                if remember_me:
                    session.permanent = True
                    # Usar configuração PERMANENT_SESSION_LIFETIME
                    session['expires_at'] = (datetime.now(timezone.utc) + app.config['PERMANENT_SESSION_LIFETIME']).timestamp()
                    session['sliding_expiry'] = False  # Não renovar automaticamente
                else:
                    session.permanent = False
                    # Usar configuração TEMPORARY_SESSION_LIFETIME com renovação automática (sliding)
                    session['expires_at'] = (datetime.now(timezone.utc) + app.config['TEMPORARY_SESSION_LIFETIME']).timestamp()
                    session['sliding_expiry'] = True  # Renovar a cada pedido

                LoginLog.log_attempt(user.id, ip_address, True)
                
                # Verificar se o utilizador tem permissões
                if user.role == 'none':
                    flash(f'Bem-vindo, {user.name}! A sua conta foi criada com sucesso mas encontra-se à espera de validação pelo administrador.', 'info')
                else:
                    flash(f'Bem-vindo, {user.name}!', 'success')
                
                return redirect(url_for('index'))
            else:
                # Password incorreta
                LoginLog.log_attempt(user.id if user else None, ip_address, False)
                
                # Verificar se deve banir o IP
                if LoginLog.check_failed_logins(ip_address):
                    flash('Demasiadas tentativas falhadas. O seu IP foi bloqueado.', 'error')
                else:
                    flash('Email ou palavra-passe incorretos.', 'error')
                
                return render_template('login.html', action='login')
        
        elif action == 'register':
            email = request.form.get('email', '').strip()
            
            if not email:
                flash('Email é obrigatório.', 'error')
                return render_template('login.html', action='register')
            
            # Validar formato do email
            if not AddUserSecurityCheck.validate_email_format(email):
                flash('Formato de email inválido.', 'error')
                return render_template('login.html', action='register')
            
            # Verificar se já existe
            if User.query.filter_by(_email=email).first() or PreUser.query.filter_by(_email=email).first():
                flash('Já existe um utilizador com este email.', 'error')
                return render_template('login.html', action='register')
            
            try:
                # Gerar código de verificação
                code = PreUser.create_random_code()
                
                # Enviar email de verificação de forma assíncrona
                from tasks import send_verification_email
                
                # Preparar configurações do app para a tarefa
                app_config = {
                    'MAIL_SERVER': app.config['MAIL_SERVER'],
                    'MAIL_PORT': app.config['MAIL_PORT'],
                    'MAIL_USE_TLS': app.config['MAIL_USE_TLS'],
                    'MAIL_USE_SSL': app.config['MAIL_USE_SSL'],
                    'MAIL_USERNAME': app.config['MAIL_USERNAME'],
                    'MAIL_PASSWORD': app.config['MAIL_PASSWORD'],
                    'MAIL_DEFAULT_SENDER': app.config['MAIL_DEFAULT_SENDER'],
                    'TEMPLATES_AUTO_RELOAD': True
                }
                
                # Enqueue da tarefa
                job = email_queue.enqueue(
                    send_verification_email,
                    app_config,
                    email,
                    code,
                    request.url_root,
                    job_timeout=300
                )
                
                # Criar pré-utilizador imediatamente (o email será enviado em background)
                pre_user = PreUser()
                pre_user._email = email  # Set directly to avoid validation during construction
                pre_user.code = code
                pre_user.reason = 'Registo de novo utilizador'
                pre_user.email_job_id = job.get_id()  # Armazenar ID da tarefa para tracking
                db.session.add(pre_user)
                db.session.commit()
                
                flash('Pedido de registo processado. O email de verificação está a ser enviado em background. Verifique a sua caixa de entrada em alguns minutos.', 'info')
                return render_template('login.html', action='verify', email=email, job_id=job.get_id())
            
            except Exception as e:
                db.session.rollback()
                print(f"Erro ao processar registo: {e}")
                print(f"Traceback completo: {traceback.format_exc()}")
                flash('Erro ao processar registo. Tente novamente.', 'error')
                return render_template('login.html', action='register')
        
        elif action == 'verify':
            email = request.form.get('email', '').strip()
            name = request.form.get('name', '').strip()
            password = request.form.get('password', '')
            confirm_password = request.form.get('confirm_password', '')
            code = request.form.get('verification_code', '').strip()
            
            if not all([email, name, password, confirm_password, code]):
                flash('Todos os campos são obrigatórios.', 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Validar formato do email
            if not AddUserSecurityCheck.validate_email_format(email):
                flash('Formato de email inválido.', 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Validar se as passwords coincidem
            if password != confirm_password:
                flash('As passwords não coincidem.', 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Validar força da password
            is_strong, password_error = AddUserSecurityCheck.validate_password_strength(password)
            if not is_strong:
                flash(password_error, 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Validar formato do código (6 caracteres alfanuméricos)
            if len(code) != 6 or not code.isalnum():
                flash('Código de verificação deve ter 6 caracteres alfanuméricos.', 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Verificar código
            pre_user = PreUser.query.filter_by(_email=email, code=code).first()
            if not pre_user:
                flash('Código de verificação inválido.', 'error')
                return render_template('login.html', action='verify', email=email)
            
            # Verificar se não expirou (10 minutos)
            now_utc = datetime.now(timezone.utc)
            pre_user_date = pre_user.date
            if pre_user_date.tzinfo is None:
                pre_user_date = pre_user_date.replace(tzinfo=timezone.utc)
            if now_utc - pre_user_date > timedelta(minutes=10):
                flash('Código de verificação expirado. Solicite um novo registo.', 'error')
                db.session.delete(pre_user)
                db.session.commit()
                return render_template('login.html', action='register')
            
            try:
                # Verificar se é o primeiro utilizador (será administrador)
                is_first_user = User.query.count() == 0
                
                # Criar utilizador final
                user = User()
                user._email = email  # Set directly to avoid validation during construction
                user.name = name  # Agora o nome vem do formulário de verificação
                user.password = password  # Definir password imediatamente
                user.role = 'admin' if is_first_user else 'none'  # Primeiro user é admin
                user.is_verified = True
                
                db.session.add(user)
                db.session.delete(pre_user)  # Remover pré-utilizador
                db.session.commit()
                
                if is_first_user:
                    flash('Conta de administrador criada com sucesso!', 'success')
                else:
                    flash('Conta criada com sucesso!', 'success')
                
                # Fazer login automático
                session['user_id'] = user.id
                session['user_email'] = user.email
                session['user_name'] = user.name
                session['user_role'] = user.role
                
                return redirect(url_for('index'))
            
            except Exception as e:
                db.session.rollback()
                flash('Erro ao criar conta. Tente novamente.', 'error')
                return render_template('login.html', action='verify', email=email)
        
        elif action == 'forgot_password':
            email = request.form.get('email', '').strip()
            
            if not email:
                flash('Email é obrigatório.', 'error')
                return render_template('login.html', action='forgot_password')
            
            # Validar formato do email
            if not AddUserSecurityCheck.validate_email_format(email):
                flash('Formato de email inválido.', 'error')
                return render_template('login.html', action='forgot_password')
            
            # Verificar se o utilizador existe
            user = User.query.filter_by(_email=email).first()
            if not user:
                # Por segurança, não revelar se o email existe ou não
                flash('Se o email existir na nossa base de dados, receberá instruções para recuperar a sua password.', 'info')
                return render_template('login.html', action='login')

            # Verificar se já existe um pedido de recuperação recente
            pre_user_reset = PreUser.query.filter_by(_email=email).first()
            now_utc = datetime.now(timezone.utc)
            if pre_user_reset:
                # Garantir que pre_user_reset.date é timezone-aware
                if pre_user_reset.date.tzinfo is None:
                    pre_user_reset_date = pre_user_reset.date.replace(tzinfo=timezone.utc)
                else:
                    pre_user_reset_date = pre_user_reset.date
                if (now_utc - pre_user_reset_date <= timedelta(minutes=10)):
                    flash('Já existe um pedido de recuperação recente. Por favor, verifique o seu email ou aguarde alguns minutos antes de pedir novamente.', 'info')
                    return render_template('login.html', action='forgot_password')

            # Remover entradas PreUser anteriores para este email (password reset anteriores)
            PreUser.query.filter_by(_email=email).delete()

            # Gerar código de recuperação (igual ao de verificação)
            reset_code = PreUser.create_random_code()

            # Enviar email de recuperação de forma assíncrona
            from tasks import send_password_reset_email
            
            # Preparar configurações do app para a tarefa
            app_config = {
                'MAIL_SERVER': app.config['MAIL_SERVER'],
                'MAIL_PORT': app.config['MAIL_PORT'],
                'MAIL_USE_TLS': app.config['MAIL_USE_TLS'],
                'MAIL_USE_SSL': app.config['MAIL_USE_SSL'],
                'MAIL_USERNAME': app.config['MAIL_USERNAME'],
                'MAIL_PASSWORD': app.config['MAIL_PASSWORD'],
                'MAIL_DEFAULT_SENDER': app.config['MAIL_DEFAULT_SENDER'],
                'TEMPLATES_AUTO_RELOAD': True
            }
            
            # Enqueue da tarefa
            job = email_queue.enqueue(
                send_password_reset_email,
                app_config,
                email,
                reset_code,
                request.url_root,
                job_timeout=300
            )

            # Criar entrada PreUser imediatamente (o email será enviado em background)
            pre_user_reset = PreUser()
            pre_user_reset._email = email  # Set directly to avoid validation during construction
            pre_user_reset.code = reset_code
            pre_user_reset.reason = 'Recuperação de password'
            pre_user_reset.email_job_id = job.get_id()  # Armazenar ID da tarefa para tracking
            db.session.add(pre_user_reset)
            db.session.commit()

            flash('Pedido de recuperação processado. As instruções estão a ser enviadas por email em background. Use o link "Já tenho código de recuperação" para introduzir o código quando o receber.', 'info')
            return render_template('login.html', action='reset_password', email=email, job_id=job.get_id())
                
        elif action == 'reset_password':
            email = request.form.get('email', '').strip()
            code = request.form.get('verification_code', '').strip()
            new_password = request.form.get('new_password', '')
            confirm_password = request.form.get('confirm_password', '')

            print(email, code, new_password, confirm_password)

            if not all([email, code, new_password, confirm_password]):
                flash('Todos os campos são obrigatórios.', 'error')
                return render_template('login.html', action='reset_password', email=email)
            
            # Validar formato do email
            if not AddUserSecurityCheck.validate_email_format(email):
                flash('Formato de email inválido.', 'error')
                return render_template('login.html', action='reset_password', email=email)

            # Validar se as passwords coincidem
            if new_password != confirm_password:
                flash('As passwords não coincidem.', 'error')
                return render_template('login.html', action='reset_password', email=email)
            
            # Validar se a password é forte
            is_strong, password_error = AddUserSecurityCheck.validate_password_strength(new_password)
            if not is_strong:
                flash(password_error, 'error')
                return render_template('login.html', action='reset_password', email=email)
            
            # Validar formato do código (6 caracteres alfanuméricos)
            if len(code) != 6 or not code.isalnum():
                flash('Código de recuperação deve ter 6 caracteres alfanuméricos.', 'error')
                return render_template('login.html', action='reset_password', email=email)
            
            # Verificar código na tabela PreUser
            pre_user_reset = PreUser.query.filter_by(_email=email, code=code).first()
            if not pre_user_reset:
                flash('Código de recuperação inválido.', 'error')
                return render_template('login.html', action='reset_password', email=email)
            
            # Verificar se não expirou (10 minutos)
            now_utc = datetime.now(timezone.utc)
            pre_user_reset_date = pre_user_reset.date
            if pre_user_reset_date.tzinfo is None:
                pre_user_reset_date = pre_user_reset_date.replace(tzinfo=timezone.utc)
            if now_utc - pre_user_reset_date > timedelta(minutes=10):
                flash('Código de recuperação expirado. Solicite uma nova recuperação.', 'error')
                db.session.delete(pre_user_reset)
                db.session.commit()
                return render_template('login.html', action='forgot_password')
            
            # Verificar se o utilizador existe
            user = User.query.filter_by(_email=email).first()
            if not user:
                flash('Utilizador não encontrado.', 'error')
                return render_template('login.html', action='login')
            
            try:
                # Atualizar password do utilizador
                user.password = new_password

                # Remover entrada PreUser (código foi usado)
                if pre_user_reset:
                    db.session.delete(pre_user_reset)
                    db.session.commit()

                flash('Password alterada com sucesso! Pode agora fazer login.', 'success')
                return redirect(url_for('login'))
            
            except Exception as e:
                db.session.rollback()
                flash('Erro ao alterar password. Tente novamente.', 'error')
                return render_template('login.html', action='reset_password', email=email)
    
    # GET request
    action = request.args.get('action', 'login')
    email = request.args.get('email', '')
    return render_template('login.html', action=action, email=email)
# End def login


@app.route('/logout')
def logout():
    # Obter o ID da sessão antes de limpar
    session_id = session.get('_id')  # Flask-Session internal ID
    
    # Limpar sessão do Flask
    session.clear()
    
    # Tentar remover sessão do Redis diretamente
    if session_id:
        try:
            redis_client = app.config['SESSION_REDIS']
            session_key = f"{app.config.get('SESSION_KEY_PREFIX', 'session:')}{session_id}"
            redis_client.delete(session_key)
        except Exception as e:
            print(f"Erro ao remover sessão do Redis: {e}")
    
    flash('Sessão terminada com sucesso.', 'success')
    return redirect(url_for('login'))
# End def logout


@app.route('/turmas/')
@required_login
@required_role('viewer')
def turmas():
    # Verificar se é uma chamada API
    is_api = request.args.get('api') == 'true'
    
    # Verificar se há mensagem de filtro salvo
    filter_saved = request.args.get('filter_saved')
    filter_message = request.args.get('filter_message')
    
    if filter_saved and filter_message:
        try:
            decoded_message = filter_message  # Já vem decodificado do JavaScript
            flash(decoded_message, 'success')
        except:
            flash('Filtros salvos com sucesso!', 'success')
    
    # Obter todas as turmas da base de dados ordenadas por ID
    turmas_query = Turma.query.order_by(Turma.id).all()
    
    # Se não há turmas e é chamada API, retornar lista vazia
    if len(turmas_query) == 0 and is_api:
        return {'turmas': []}
    
    # Se for chamada API, retornar apenas lista de nomes
    if is_api:
        turmas_list = [turma.nome for turma in turmas_query]
        return {'turmas': turmas_list}
    
    # Para chamada normal, obter informações detalhadas
    turmas_data = []
    total_alunos_geral = 0
    total_fotos_geral = 0
    
    for turma in turmas_query:
        total_alunos = len(turma.alunos)
        alunos_com_foto = sum(1 for aluno in turma.alunos if aluno.foto_tirada)
        
        # Somar aos totais gerais
        total_alunos_geral += total_alunos
        total_fotos_geral += alunos_com_foto
        
        turmas_data.append({
            'id': turma.id,
            'nome': turma.nome,
            'nome_seguro': turma.nome_seguro,
            'nome_professor': turma.nome_professor,
            'email_professor': turma.email_professor,
            'total_alunos': total_alunos,
            'alunos_com_foto': alunos_com_foto,
            'last_updated': turma.last_updated
        })
    
    return render_template('turmas.html', 
                         turmas=turmas_data, 
                         total_alunos_geral=total_alunos_geral,
                         total_fotos_geral=total_fotos_geral)
# End def turmas

# === ROTAS PARA FILTROS DE TURMAS === 
# Adicionadas para permitir persistência de filtros entre sessões


@app.route('/turmas/filters', methods=['GET'])
@csrf.exempt  # API endpoint para JSON, isento de CSRF
@required_login
@required_role('viewer')
def get_turma_filters():
    """Retorna os filtros de turmas salvos na sessão do usuário"""
    try:
        filters = session.get('turma_filters', {
            'showAll': True,
            'visibleTurmas': []
        })
        
        #print(f"Filtros carregados da sessão: {filters}")
        
        # Garantir que visibleTurmas é uma lista
        if 'visibleTurmas' in filters:
            if isinstance(filters['visibleTurmas'], set):
                filters['visibleTurmas'] = list(filters['visibleTurmas'])
        
        return jsonify(filters)
        
    except Exception as e:
        #print(f"Erro ao carregar filtros: {str(e)}")
        return jsonify({'error': f'Erro ao carregar filtros: {str(e)}'}), 500


@app.route('/turmas/filters', methods=['POST'])
@csrf.exempt  # API endpoint para JSON, isento de CSRF
@required_login
@required_role('viewer')
def save_turma_filters():
    """Salva os filtros de turmas na sessão do usuário"""
    try:
        # Debug: log do request
        #print(f"Content-Type: {request.content_type}")
        #print(f"Raw data: {request.get_data()}")
        
        data = request.get_json()
        #print(f"Parsed JSON: {data}")
        
        if not data:
            #print("Erro: Dados não fornecidos")
            return jsonify({'error': 'Dados não fornecidos'}), 400
        
        # Converter Set para lista se necessário (JSON não suporta Set)
        if 'visibleTurmas' in data:
            if isinstance(data['visibleTurmas'], set):
                data['visibleTurmas'] = list(data['visibleTurmas'])
            elif hasattr(data['visibleTurmas'], '__iter__') and not isinstance(data['visibleTurmas'], str):
                data['visibleTurmas'] = list(data['visibleTurmas'])
        
        # Validar estrutura dos dados
        filters = {
            'showAll': data.get('showAll', True),
            'visibleTurmas': data.get('visibleTurmas', [])
        }

        #print(f"Filtros a salvar: {filters}")

        # Salvar na sessão
        session['turma_filters'] = filters
        session.permanent = True  # Tornar sessão permanente
        
        print("Filtros salvos com sucesso")
        return jsonify({'success': True, 'message': 'Filtros salvos com sucesso'})
        
    except Exception as e:
        print(f"Erro ao salvar filtros: {str(e)}")
        return jsonify({'error': f'Erro ao salvar filtros: {str(e)}'}), 500


@app.route('/turma/<nome_seguro>/')
@no_cache # Decorator para evitar cache das fotografias
@required_login
@required_role('viewer')
def turma(nome_seguro):
    # Buscar turma na base de dados usando nome_seguro
    turma_obj = Turma.query.filter_by(nome_seguro=nome_seguro).first()
    if not turma_obj:
        return redirect(url_for('settings'))
    
    # Obter utilizador atual
    current_user = get_current_user()
    
    # Obter alunos da turma ordenados por número (nulls last) e depois por nome
    alunos = []
    fotos_existentes = 0
    
    # Ordenar: primeiro por número (None/NULL por último), depois por nome
    alunos_ordenados = sorted(
        turma_obj.alunos, 
        key=lambda x: (x.numero is None, x.numero if x.numero is not None else 0, x.nome)
    )
    
    for aluno in alunos_ordenados:
        alunos.append({
            'id': aluno.id,
            'processo': aluno.processo,
            'nome': aluno.nome,
            'numero': aluno.numero,
            'email': aluno.email or '',
            'notes': aluno.notes or '',
            'autorizacao': aluno.autorizacao,
            'turma': turma_obj.nome,
            'foto_tirada': aluno.foto_tirada,
            'foto_existe': aluno.foto_existe
        })
        
        if aluno.foto_tirada:
            fotos_existentes += 1
    
    return render_template('turma.html', 
                         alunos=alunos, 
                         fotos_existentes=fotos_existentes,
                         current_user=current_user,
                         current_turma=turma_obj)
# End def turma


@app.route('/turma/', methods=['POST'])
@required_login
@required_role('admin')
def turma_crud():
    action = request.form.get('action', '').strip()
    
    if action == 'crud_turma_add_new':
        nome = request.form.get('nome', '').strip()
        nome_professor = request.form.get('nome_professor', '').strip()
        email_professor = request.form.get('email_professor', '').strip()
        
        if not nome:
            flash('Nome da turma é obrigatório!', 'error')
            return redirect(url_for('turmas'))
        
        # Verificar se a turma já existe
        turma_existente = Turma.query.filter_by(nome=nome).first()
        if turma_existente:
            flash(f'Já existe uma turma com o nome "{nome}"!', 'error')
            return redirect(url_for('turmas'))
        
        # Criar nova turma usando o método seguro
        try:
            nova_turma = Turma.create_safe(nome)
            nova_turma.nome_professor = nome_professor
            nova_turma.email_professor = email_professor
            db.session.commit()
            flash(f'Turma "{nome}" criada com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao criar turma. Tente novamente.', 'error')
        
        return redirect(url_for('turmas'))
    
    elif action == 'crud_turma_edit':
        turma_id = request.form.get('turma_id', '').strip()
        nome = request.form.get('nome', '').strip()
        nome_professor = request.form.get('nome_professor', '').strip()
        email_professor = request.form.get('email_professor', '').strip()
        
        if not turma_id or not nome:
            flash('Dados obrigatórios não fornecidos!', 'error')
            return redirect(url_for('turmas'))
        
        # Buscar turma
        turma = db.session.get(Turma, turma_id)
        if not turma:
            flash('Turma não encontrada!', 'error')
            return redirect(url_for('turmas'))
        
        # Verificar se o novo nome já existe (exceto para a própria turma)
        if nome != turma.nome:
            turma_existente = Turma.query.filter_by(nome=nome).first()
            if turma_existente:
                flash(f'Já existe outra turma com o nome "{nome}"!', 'error')
                return redirect(url_for('turmas'))
        
        nome_antigo = turma.nome
        
        try:
            # Atualizar nome da turma usando método seguro
            turma.update_nome(nome)
            turma.nome_professor = nome_professor
            turma.email_professor = email_professor
            db.session.commit()
            flash(f'Turma editada com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao editar turma. Tente novamente.', 'error')
        
        return redirect(url_for('turmas'))
    
    elif action == 'crud_turma_delete':
        turma_id = request.form.get('turma_id', '').strip()
        
        if not turma_id:
            flash('ID da turma não fornecido!', 'error')
            return redirect(url_for('turmas'))
        
        # Buscar turma
        turma = db.session.get(Turma, turma_id)
        if not turma:
            flash('Turma não encontrada!', 'error')
            return redirect(url_for('turmas'))
        
        nome_turma = turma.nome
        
        try:
            # Remover diretórios de fotos e thumbnails usando método seguro
            turma.delete_directories()
            
            # Remover turma da base de dados (alunos são removidos automaticamente pelo cascade)
            db.session.delete(turma)
            db.session.commit()
            flash(f'Turma "{nome_turma}" removida com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao remover turma. Tente novamente.', 'error')
        
        return redirect(url_for('turmas'))
    
    # Se a ação não for reconhecida, redirecionar
    flash('Ação não reconhecida!', 'error')
    return redirect(url_for('turmas'))
# End def turma_crud (POST only)


@app.route('/student/', methods=['POST'])
@required_login
@required_role('editor')
def student_crud():
    action = request.form.get('action', '').strip()
    turma = request.form.get('turma', '').strip()
    
    if action == 'crud_student_add_new':
        nome = request.form.get('nome', '').strip()
        processo = request.form.get('processo', '').strip()
        numero = request.form.get('numero', '').strip()
        email = request.form.get('email', '').strip()
        notes = request.form.get('notes', '').strip()
        autorizacao = request.form.get('autorizacao') == 'on'  # Checkbox (default True se não marcado)
        
        if not nome:
            flash('Nome é obrigatório!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Validar e sanitizar processo
        processo_validado, erro_processo = Aluno.validate_and_sanitize_processo(processo)
        if erro_processo:
            flash(erro_processo, 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Converter número para inteiro se fornecido
        numero_int = None
        if numero:
            try:
                numero_int = int(numero)
            except ValueError:
                flash('O número na turma deve ser um número válido!', 'error')
                return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Verificar se a turma existe
        turma_obj = Turma.query.filter_by(nome=turma).first()
        if not turma_obj:
            flash('Turma não encontrada!', 'error')
            return redirect(url_for('turmas'))
        
        # Verificar se o processo já existe globalmente
        aluno_existente = Aluno.query.filter_by(processo=processo_validado).first()
        if aluno_existente:
            flash(f'Já existe um aluno com o número {processo_validado} na turma "{aluno_existente.turma.nome}"!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Criar novo aluno
        novo_aluno = Aluno(
            nome=nome,
            processo=processo_validado,
            numero=numero_int,
            email=email,
            notes=notes,
            autorizacao=autorizacao,
            turma_id=turma_obj.id,
            foto_tirada=False
        )
        
        try:
            db.session.add(novo_aluno)
            db.session.commit()
            flash(f'Aluno {nome} adicionado com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao adicionar aluno. Tente novamente.', 'error')
        
        return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
    
    elif action == 'crud_student_edit':
        aluno_id = request.form.get('aluno_id', '').strip()
        nome = request.form.get('nome', '').strip()
        processo = request.form.get('processo', '').strip()
        numero = request.form.get('numero', '').strip()
        email = request.form.get('email', '').strip()
        notes = request.form.get('notes', '').strip()
        autorizacao = request.form.get('autorizacao') == 'on'  # Checkbox

        if not aluno_id or not nome:
            flash('Dados obrigatórios não fornecidos!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Validar e sanitizar processo
        processo_validado, erro_processo = Aluno.validate_and_sanitize_processo(processo)
        if erro_processo:
            flash(erro_processo, 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Converter número para inteiro se fornecido
        numero_int = None
        if numero:
            try:
                numero_int = int(numero)
            except ValueError:
                flash('O número na turma deve ser um número válido!', 'error')
                return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Buscar aluno
        aluno = db.session.get(Aluno, aluno_id)
        if not aluno:
            flash('Aluno não encontrado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Verificar se o novo processo já existe globalmente (exceto para o próprio aluno)
        if processo_validado != aluno.processo:
            aluno_existente = Aluno.query.filter_by(processo=processo_validado).first()
            if aluno_existente and aluno_existente.id != aluno.id:
                flash(f'Já existe outro aluno com o número {processo_validado} na turma "{aluno_existente.turma.nome}"!', 'error')
                return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Guardar processo antigo para renomeação de arquivos
        processo_antigo = aluno.processo
        processo_mudou = (processo_validado != processo_antigo)
        
        # Se o processo mudou e o aluno tem foto, renomear arquivos
        if processo_mudou and aluno.foto_tirada:
            foto_dir = aluno.turma.get_foto_directory()
            thumb_dir = aluno.turma.get_thumb_directory()
            
            # Caminhos dos arquivos antigos
            old_photo_path = os.path.join(foto_dir, f'{processo_antigo}.jpg')
            old_thumb_path = os.path.join(thumb_dir, f'{processo_antigo}.jpg')
            
            # Caminhos dos arquivos novos
            new_photo_path = os.path.join(foto_dir, f'{processo_validado}.jpg')
            new_thumb_path = os.path.join(thumb_dir, f'{processo_validado}.jpg')
            
            # Renomear foto original se existir
            if os.path.exists(old_photo_path):
                try:
                    os.rename(old_photo_path, new_photo_path)
                except Exception as e:
                    print(f"Erro ao renomear foto original de {processo_antigo} para {processo_validado}: {e}")
                    flash('Erro ao renomear foto original. Tente novamente.', 'error')
                    return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
            
            # Renomear thumbnail se existir
            if os.path.exists(old_thumb_path):
                try:
                    os.rename(old_thumb_path, new_thumb_path)
                except Exception as e:
                    print(f"Erro ao renomear thumbnail de {processo_antigo} para {processo_validado}: {e}")
                    # Se houve erro na thumbnail mas a foto original foi renomeada, tentar desfazer
                    if os.path.exists(new_photo_path):
                        try:
                            os.rename(new_photo_path, old_photo_path)
                        except:
                            pass
                    flash('Erro ao renomear thumbnail. Tente novamente.', 'error')
                    return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Atualizar dados na base de dados
        try:
            aluno.nome = nome
            aluno.processo = processo_validado
            aluno.numero = numero_int
            aluno.email = email
            aluno.notes = notes
            aluno.autorizacao = autorizacao
            
            # Se o processo mudou e há foto, atualizar timestamp da turma
            if processo_mudou and aluno.foto_tirada:
                aluno.turma.update_last_modified()
                
            db.session.commit()
            
            if processo_mudou and aluno.foto_tirada:
                flash(f'Aluno {nome} editado com sucesso! Fotos renomeadas de {processo_antigo} para {processo_validado}.', 'success')
            else:
                flash(f'Aluno {nome} editado com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao editar aluno na base de dados. Tente novamente.', 'error')
            print(f"Erro na edição do aluno: {e}")
        
        return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
    
    elif action == 'crud_student_transfer':
        aluno_id = request.form.get('aluno_id', '').strip()
        turma_destino = request.form.get('turma_destino', '').strip()
        
        if not aluno_id or not turma_destino:
            flash('Dados obrigatórios não fornecidos!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Buscar aluno
        aluno = db.session.get(Aluno, aluno_id)
        if not aluno:
            flash('Aluno não encontrado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Buscar turma de destino
        turma_destino_obj = Turma.query.filter_by(nome=turma_destino).first()
        if not turma_destino_obj:
            flash('Turma de destino não encontrada!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        nome_aluno = aluno.nome
        processo_aluno = aluno.processo
        turma_origem = aluno.turma.nome
        turma_origem_obj = aluno.turma  # Guardar referência para atualizar timestamp
        
        try:
            # Mover arquivos de foto se existirem
            old_photo_path = os.path.join(aluno.turma.get_foto_directory(), f'{processo_aluno}.jpg')
            old_thumb_path = os.path.join(aluno.turma.get_thumb_directory(), f'{processo_aluno}.jpg')
            new_photo_path = os.path.join(turma_destino_obj.get_foto_directory(), f'{processo_aluno}.jpg')
            new_thumb_path = os.path.join(turma_destino_obj.get_thumb_directory(), f'{processo_aluno}.jpg')
            
            # Criar diretórios de destino se não existirem
            safe_makedirs(os.path.dirname(new_photo_path))
            safe_makedirs(os.path.dirname(new_thumb_path))
            
            # Mover arquivos se existirem
            import shutil
            foto_movida = False
            if os.path.exists(old_photo_path):
                shutil.move(old_photo_path, new_photo_path)
                foto_movida = True
            if os.path.exists(old_thumb_path):
                shutil.move(old_thumb_path, new_thumb_path)
            
            # Atualizar turma do aluno na base de dados
            aluno.turma_id = turma_destino_obj.id
            
            # Atualizar timestamps das turmas se houve movimento de fotos
            if foto_movida and aluno.foto_tirada:
                turma_origem_obj.update_last_modified()  # Turma que perdeu a foto
                turma_destino_obj.update_last_modified()  # Turma que ganhou a foto
            
            db.session.commit()
            flash(f'Aluno {nome_aluno} transferido com sucesso para a turma {turma_destino}!', 'success')
            
        except Exception as e:
            db.session.rollback()
            flash('Erro ao transferir aluno. Tente novamente.', 'error')
            print(f"Erro na transferência: {e}")
        
        return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
    
    elif action == 'crud_student_delete':
        aluno_id = request.form.get('aluno_id', '').strip()
        
        if not aluno_id:
            flash('ID do aluno não fornecido!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        # Buscar aluno
        aluno = db.session.get(Aluno, aluno_id)
        if not aluno:
            flash('Aluno não encontrado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
        
        nome_aluno = aluno.nome
        processo_aluno = aluno.processo
        turma_obj = aluno.turma  # Guardar referência da turma
        tinha_foto = aluno.foto_tirada  # Verificar se tinha foto
        
        try:
            # Remover arquivos de foto se existirem usando métodos seguros
            foto_dir = aluno.turma.get_foto_directory()
            thumb_dir = aluno.turma.get_thumb_directory()
            photo_path = os.path.join(foto_dir, f'{processo_aluno}.jpg')
            thumb_path = os.path.join(thumb_dir, f'{processo_aluno}.jpg')
            
            if os.path.exists(photo_path):
                os.remove(photo_path)
            if os.path.exists(thumb_path):
                os.remove(thumb_path)
            
            # Atualizar timestamp da turma se havia foto
            if tinha_foto:
                turma_obj.update_last_modified()
            
            # Remover aluno da base de dados
            db.session.delete(aluno)
            db.session.commit()
            flash(f'Aluno {nome_aluno} removido com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao remover aluno. Tente novamente.', 'error')
        
        return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
    

    elif action == 'crud_student_remove_photo':
        aluno_id = request.form.get('aluno_id', '').strip()
        if not aluno_id:
            flash('ID do aluno não fornecido!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        # Buscar aluno
        aluno = db.session.get(Aluno, aluno_id)
        if not aluno:
            flash('Aluno não encontrado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        # Verificar se o aluno tem foto
        if not aluno.foto_tirada:
            flash('Este aluno não tem foto para remover!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        nome_aluno = aluno.nome
        processo_aluno = aluno.processo
        turma_obj = aluno.turma

        try:
            # Remover arquivos de foto se existirem usando métodos seguros
            foto_dir = turma_obj.get_foto_directory()
            thumb_dir = turma_obj.get_thumb_directory()
            photo_path = os.path.join(foto_dir, f'{processo_aluno}.jpg')
            thumb_path = os.path.join(thumb_dir, f'{processo_aluno}.jpg')

            if os.path.exists(photo_path):
                os.remove(photo_path)
            if os.path.exists(thumb_path):
                os.remove(thumb_path)

            # Atualizar flag de foto na base de dados
            aluno.foto_tirada = False
            aluno.foto_existe = False  # Sempre que foto é removida, foto_existe também
            aluno.turma.update_last_modified()  # Atualizar timestamp da turma
            db.session.commit()

            # Verificar se a turma ficou sem fotos tiradas
            fotos_tiradas = sum(1 for a in turma_obj.alunos if a.foto_tirada)
            if fotos_tiradas == 0:
                # Remover as pastas da turma se estiverem vazias
                try:
                    if os.path.isdir(foto_dir) and len(os.listdir(foto_dir)) == 0:
                        os.rmdir(foto_dir)
                    if os.path.isdir(thumb_dir) and len(os.listdir(thumb_dir)) == 0:
                        os.rmdir(thumb_dir)
                except Exception as e:
                    print(f"Erro ao remover pastas vazias da turma: {e}")

            flash(f'Foto do aluno {nome_aluno} removida com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao remover foto. Tente novamente.', 'error')
            print(f"Erro ao remover foto: {e}")

        return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
    
    elif action == 'crud_student_manual_upload_photo':
        aluno_id = request.form.get('aluno_id', '').strip()
        if not aluno_id:
            flash('ID do aluno não fornecido!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        # Buscar aluno
        aluno = db.session.get(Aluno, aluno_id)
        if not aluno:
            flash('Aluno não encontrado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        # Verificar se arquivo foi enviado
        if 'foto' not in request.files:
            flash('Nenhum arquivo foi selecionado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        foto_file = request.files['foto']
        if foto_file.filename == '':
            flash('Nenhum arquivo foi selecionado!', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        # Verificar se é um arquivo de imagem válido
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}
        if not ('.' in foto_file.filename and foto_file.filename.rsplit('.', 1)[1].lower() in allowed_extensions):
            flash('Tipo de arquivo não permitido! Use PNG, JPG, JPEG, GIF, BMP ou WEBP.', 'error')
            return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

        nome_aluno = aluno.nome
        processo_aluno = aluno.processo
        turma_obj = aluno.turma

        try:
            # Criar diretórios se não existirem
            foto_dir = turma_obj.get_foto_directory()
            thumb_dir = turma_obj.get_thumb_directory()
            
            if not safe_makedirs(foto_dir):
                flash('Erro ao criar diretório para fotos.', 'error')
                return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
            
            if not safe_makedirs(thumb_dir):
                flash('Erro ao criar diretório para thumbnails.', 'error')
                return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

            # Caminhos dos arquivos
            photo_path = os.path.join(foto_dir, f'{processo_aluno}.jpg')
            thumb_path = os.path.join(thumb_dir, f'{processo_aluno}.jpg')

            # Ler e processar a imagem
            foto_data = foto_file.read()
            image = cv2.imdecode(np.frombuffer(foto_data, np.uint8), cv2.IMREAD_COLOR)
            
            if image is None:
                flash('Arquivo de imagem inválido ou corrompido!', 'error')
                return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))

            # Salvar foto original
            cv2.imwrite(photo_path, image, [cv2.IMWRITE_JPEG_QUALITY, 95])

            # Criar thumbnail (crop quadrado e redimensionar)
            height, width, _ = image.shape
            min_dim = min(height, width)
            start_x = (width - min_dim) // 2
            start_y = (height - min_dim) // 2
            cropped_image = image[start_y:start_y + min_dim, start_x:start_x + min_dim]
            thumbnail = cv2.resize(cropped_image, (250, 250))
            cv2.imwrite(thumb_path, thumbnail, [cv2.IMWRITE_JPEG_QUALITY, 50])

            # Atualizar flags na base de dados
            aluno.foto_tirada = True
            aluno.foto_existe = True
            aluno.turma.update_last_modified()  # Atualizar timestamp da turma
            db.session.commit()

            flash(f'Foto do aluno {nome_aluno} enviada com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash('Erro ao processar foto. Tente novamente.', 'error')
            print(f"Erro ao processar upload manual de foto: {e}")

        return redirect(url_for('turma', nome_seguro=Turma.get_nome_seguro_by_nome(turma)))
    
    # Se a ação não for reconhecida, redirecionar
    flash('Ação não reconhecida!', 'error')
    return redirect(url_for('turmas'))
# End def student_crud (POST only)


@app.route('/upload/photo/<nome_seguro>/<processo>', methods=['POST'])
@csrf.exempt  # API endpoint que recebe JSON, não formulários HTML
@required_login
@required_role('editor')
def upload_photo(nome_seguro, processo):
    data = request.get_json()
    if not data or 'image' not in data:
        return "Dados inválidos.", 400

    # Buscar a turma usando nome_seguro
    turma_obj = Turma.query.filter_by(nome_seguro=nome_seguro).first()
    if not turma_obj:
        return "Turma não encontrada.", 404

    image_data = data['image'].split(',')[1]
    foto_dir = turma_obj.get_foto_directory()
    thumb_dir = turma_obj.get_thumb_directory()
    photo_path = os.path.join(foto_dir, f'{processo}.jpg')
    
    # Criar diretório da foto de forma segura
    if not safe_makedirs(os.path.dirname(photo_path)):
        return "Erro ao criar diretório para fotos.", 500

    try:
        with open(photo_path, 'wb') as photo_file:
            photo_file.write(BytesIO(base64.b64decode(image_data)).getvalue())
    except Exception as e:
        return f"Erro ao salvar foto: {e}", 500

    thumb_path = os.path.join(thumb_dir, f'{processo}.jpg')
    
    # Criar diretório da thumbnail de forma segura
    if not safe_makedirs(os.path.dirname(thumb_path)):
        return "Erro ao criar diretório para thumbnails.", 500

    with open(photo_path, 'rb') as photo_file:
        image = cv2.imdecode(np.frombuffer(photo_file.read(), np.uint8), cv2.IMREAD_COLOR)
        height, width, _ = image.shape
        min_dim = min(height, width)
        start_x = (width - min_dim) // 2
        start_y = (height - min_dim) // 2
        cropped_image = image[start_y:start_y + min_dim, start_x:start_x + min_dim]
        thumbnail = cv2.resize(cropped_image, (250, 250))
        cv2.imwrite(photo_path, cv2.imdecode(cv2.imencode('.jpg', image, [cv2.IMWRITE_JPEG_QUALITY, 95])[1], cv2.IMREAD_COLOR))
        cv2.imwrite(thumb_path, cv2.imdecode(cv2.imencode('.jpg', thumbnail, [cv2.IMWRITE_JPEG_QUALITY, 50])[1], cv2.IMREAD_COLOR))

    # Atualizar flag de foto tirada na base de dados
    aluno_obj = Aluno.query.join(Turma).filter(
        Turma.nome_seguro == nome_seguro,
        Aluno.processo == processo
    ).first()
    
    if aluno_obj:
        aluno_obj.foto_tirada = True
        aluno_obj.foto_existe = True  # Sempre que foto é tirada, foto_existe também
        aluno_obj.turma.update_last_modified()  # Atualizar timestamp da turma
        db.session.commit()
    else:
        # Retornar erro se não conseguir atualizar aluno
        return "Erro ao atualizar informações do aluno.", 500

    return "Foto enviada com sucesso.", 200
# End def upload_photo (csrf exempt)


@app.route('/photos/<folder_name>/<processo>.jpg')
@no_cache
@required_login
@required_role('viewer')
def get_photo(folder_name, processo):
    # Determinar se é pedido de thumbnail ou foto original
    # Usar parâmetro 'size' na query string: ?size=original ou ?size=thumb
    size = request.args.get('size', 'thumb')
    is_original = size == 'original'
    
    # Buscar a turma usando nome_seguro
    turma_obj = Turma.query.filter_by(nome_seguro=folder_name).first()
    if not turma_obj:
        return send_file(os.path.join(BASE_DIR, 'static', 'student_icon.jpg'))
    
    photo_dir = turma_obj.get_foto_directory() if is_original else turma_obj.get_thumb_directory()
    photo_path = os.path.join(photo_dir, f'{processo}.jpg')
    
    #print(photo_path)
    if not os.path.exists(photo_path):
        #print(f"Foto não encontrada: {photo_path}")
        return send_file(os.path.join(BASE_DIR, 'static', 'student_icon.jpg'))
    #print(f"Enviando foto: {photo_path}")
    return send_file(photo_path)
# End def get_photo


@app.route('/api/login', methods=['POST'])
@csrf.exempt
def api_login():
    # Obter IP address
    ip_address = request.headers.get('Cf-Connecting-Ip', None)  # Cloudflare proxy
    if not ip_address:
        ip_address = request.access_route[-1] if request.access_route else request.remote_addr
    
    # Verificar se o IP está banido
    DESATIVAR = False  # Mesmo que na rota normal
    if BannedIPs.is_banned(ip_address) and not DESATIVAR:
        return jsonify({'error': 'Your IP has been blocked due to excessive login attempts. Contact the administrator.'}), 403
    
    data = request.get_json()
    if not data or 'username' not in data or 'password' not in data:
        return jsonify({'error': 'Username and password required'}), 400
    
    email = data['username'].strip()
    password = data['password']
    remember_me = data.get('remember_me', False)
    
    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400
    
    # Validar formato do email
    if not AddUserSecurityCheck.validate_email_format(email):
        return jsonify({'error': 'Invalid email format'}), 400
    
    # Verificar utilizador
    user = User.query.filter_by(_email=email).first()
    
    if not user:
        # Log da tentativa mesmo sem utilizador para detectar ataques
        LoginLog.log_attempt(None, ip_address, False)
        
        # Verificar se deve banir o IP
        if LoginLog.check_failed_logins(ip_address):
            return jsonify({'error': 'Too many failed attempts. Your IP has been blocked.'}), 429
        else:
            return jsonify({'error': 'Invalid credentials'}), 401
    
    # Verificar se a conta está verificada
    if not user.is_verified:
        return jsonify({'error': 'Account not verified'}), 401
    
    # Verificar se o utilizador tem password definida
    if not user.password_hash:
        return jsonify({'error': 'Password not set'}), 401
    
    if user.check_password(password):
        # Login bem-sucedido
        session['user_id'] = user.id
        session['user_email'] = user.email
        session['user_role'] = user.role
        
        # Configurar duração da sessão baseado no "remember me"
        if remember_me:
            session.permanent = True
            session['expires_at'] = (datetime.now(timezone.utc) + app.config['PERMANENT_SESSION_LIFETIME']).timestamp()
            session['sliding_expiry'] = False
        else:
            session.permanent = False
            session['expires_at'] = (datetime.now(timezone.utc) + app.config['TEMPORARY_SESSION_LIFETIME']).timestamp()
            session['sliding_expiry'] = True
        
        # Log da tentativa bem-sucedida
        LoginLog.log_attempt(user.id, ip_address, True)
        
        # Gerar token de sessão válido por 1 hora
        session_token = serializer.dumps({'user_id': user.id}, salt='session_token')
        
        return jsonify({
            'success': True, 
            'message': f'Welcome, {user.name}!',
            'user': {
                'id': user.id,
                'name': user.name,
                'role': user.role
            },
            'token': session_token
        })
    else:
        # Password incorreta
        LoginLog.log_attempt(user.id, ip_address, False)
        
        # Verificar se deve banir o IP
        if LoginLog.check_failed_logins(ip_address):
            return jsonify({'error': 'Too many failed attempts. Your IP has been blocked.'}), 429
        else:
            return jsonify({'error': 'Invalid credentials'}), 401
# End function api_login


@app.route('/api/photos/', methods=['GET'])
def api_get_photo():
    token = request.args.get('token')
    processo = request.args.get('processo')
    noplaceholder = request.args.get('noplaceholder', 'false').lower() == 'true'
    # Determinar se é pedido de thumbnail ou foto original
    size = request.args.get('size', 'thumb')
    is_original = size == 'original'
    
    if not token:
        return jsonify({'error': 'Token required'}), 400
    
    if not processo:
        return jsonify({'error': 'Processo required'}), 400
    
    try:
        data = serializer.loads(token, salt='session_token', max_age=3600)
    except Exception as e:
        return jsonify({'error': 'Invalid or expired token'}), 400
    
    user_id = data['user_id']
    
    # Buscar aluno por processo
    aluno = Aluno.query.filter_by(processo=processo).first()
    if not aluno:
        if noplaceholder:
            return jsonify({'error': 'Student not found', 'processo': processo}), 404
        return send_file(os.path.join(BASE_DIR, 'static', 'student_icon.jpg'))
    
    # Obter turma e folder
    turma = aluno.turma
    if not turma:
        if noplaceholder:
            return jsonify({'error': 'Class not found for student', 'processo': processo}), 404
        return send_file(os.path.join(BASE_DIR, 'static', 'student_icon.jpg'))
    
    folder = turma.nome_seguro
    
    # Buscar a turma usando nome_seguro
    turma_obj = Turma.query.filter_by(nome_seguro=folder).first()
    if not turma_obj:
        if noplaceholder:
            return jsonify({'error': 'Class folder not found', 'processo': processo, 'folder': folder}), 404
        return send_file(os.path.join(BASE_DIR, 'static', 'student_icon.jpg'))
    
    photo_dir = turma_obj.get_foto_directory() if is_original else turma_obj.get_thumb_directory()
    photo_path = os.path.join(photo_dir, f'{processo}.jpg')
    
    if not os.path.exists(photo_path):
        if noplaceholder:
            return jsonify({
                'error': 'Photo not found', 
                'processo': processo, 
                'size': size,
                'path': photo_path
            }), 404
        return send_file(os.path.join(BASE_DIR, 'static', 'student_icon.jpg'))
    
    return send_file(photo_path)
# End function api_get_photo


@app.route('/download/<ficheiro>')
@required_login
def download(ficheiro=None):
    # Extrair nome da turma e extensão do ficheiro
    if '.' in ficheiro:
        turma, extensao = ficheiro.rsplit('.', 1)
    else:
        # Se não houver extensão, assumir que é zip
        turma = ficheiro
        extensao = 'zip'
    
    # Processar baseado na extensão
    if extensao.lower() == 'zip':
        # Verificar se é download de thumbnails
        if turma.endswith('.thumbs'):
            # Remove '.thumbs' do nome da turma
            turma_nome = turma[:-7]  # Remove '.thumbs'
            is_thumb = True
            download_filename = f'{turma_nome}.thumbs.zip'
        else:
            # Download normal das fotos
            turma_nome = turma
            is_thumb = False
            download_filename = f'{turma_nome}.zip'
        
        # Buscar a turma para usar métodos seguros
        # Primeiro tentar procurar por nome_seguro (novo formato), depois por nome (compatibilidade)
        turma_obj = Turma.query.filter_by(nome_seguro=turma_nome).first()
        if not turma_obj:
            turma_obj = Turma.query.filter_by(nome=turma_nome).first()
        
        if not turma_obj:
            return "Turma não encontrada.", 404
        
        # Para o nome do ficheiro de download, usar sempre o nome original da turma
        real_turma_nome = turma_obj.nome
        #download_filename = f'{real_turma_nome}.thumbs.zip' if is_thumb else f'{real_turma_nome}.zip'
        
        turma_dir = turma_obj.get_thumb_directory() if is_thumb else turma_obj.get_foto_directory()
        
        # Cria e envia o ZIP com as imagens
        files = [file for root, dirs, files in os.walk(turma_dir) for file in files]

        if not files:
            return "Nenhuma imagem disponível para download.", 400

        memory_file = BytesIO()
        with zipfile.ZipFile(memory_file, 'w') as zipf:
            for file in files:
                zipf.write(os.path.join(turma_dir, file), arcname=file)
        memory_file.seek(0)
        return send_file(memory_file, as_attachment=True, download_name=download_filename)
    
    elif extensao.lower() == 'docx':
        # Buscar a turma para usar métodos seguros
        # Primeiro tentar procurar por nome_seguro (novo formato), depois por nome (compatibilidade)
        turma_obj = Turma.query.filter_by(nome_seguro=turma).first()
        if not turma_obj:
            turma_obj = Turma.query.filter_by(nome=turma).first()
        
        if not turma_obj:
            return "Turma não encontrada.", 404
        
        # Cria e envia o documento DOCX com as fotos da turma
        docx_file = create_docx_with_photos(turma_obj.nome)
        
        if not docx_file:
            return "Erro ao criar documento ou nenhuma foto disponível.", 400
        
        return send_file(
            docx_file, 
            as_attachment=True, 
            download_name=f'{turma_obj.nome_seguro}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        return "Tipo de ficheiro não suportado.", 400
# End def download (docx, zip, thumbs.zip)


@app.route('/settings/')
@required_login
@required_role('admin')
def settings():
    user = get_current_user()
    # Verificar se existem dados na base de dados
    has_data = Turma.query.count() > 0
    
    # Obter todos os utilizadores para gestão
    users = User.query.all()
    preusers = PreUser.query.all()
    login_logs = LoginLog.query.order_by(LoginLog.date.desc()).limit(100).all()
    banned_ips = BannedIPs.query.order_by(BannedIPs.date.desc()).all()
    
    return render_template(
        'settings.html', 
        has_data=has_data, 
        user=user,
        users=users,
        preusers=preusers,
        login_logs=login_logs,
        banned_ips=banned_ips,
        current_user=user,
        can_upload_csv=True,
        can_system_nuke=True,
        can_manage_turmas=True
    )
# End def settings


# Rota para rescan de fotos e atualização de flag foto_tirada
@app.route('/settings/rescan_photos/', methods=['POST'])
@required_login
@required_role('admin')
def settings_rescan_photos():
    """Percorre todas as pastas de fotos e atualiza foto_existe=True para alunos com foto encontrada"""
    import os
    
    # Obter opção de marcação do formulário
    mark_option = request.form.get('mark_option', 'taken')  # Default: marcar como tiradas
    mark_as_taken = (mark_option == 'taken')
    
    updated_count = 0
    orphaned_photos_count = 0
    orphaned_thumbs_count = 0
    
    # Criar pastas "lost" se necessário
    lost_photos_dir = os.path.join(PHOTOS_DIR, 'lost')
    lost_thumbs_dir = os.path.join(THUMBS_DIR, 'lost')
    safe_makedirs(lost_photos_dir)
    safe_makedirs(lost_thumbs_dir)
    
    # Processar fotos originais
    for turma_dir in os.listdir(PHOTOS_DIR) if os.path.exists(PHOTOS_DIR) else []:
        if turma_dir == 'lost':  # Pular a pasta lost
            continue
            
        turma_path = os.path.join(PHOTOS_DIR, turma_dir)
        if os.path.isdir(turma_path):
            for photo_file in os.listdir(turma_path):
                if photo_file.endswith('.jpg'):
                    processo = photo_file[:-4]  # Remove .jpg
                    aluno = Aluno.query.filter_by(processo=processo).first()
                    if aluno:
                        aluno.foto_existe = True
                        aluno.foto_tirada = mark_as_taken  # Definir baseado na opção escolhida
                        updated_count += 1
                    else:
                        # Foto órfã - mover para pasta lost
                        source_path = os.path.join(turma_path, photo_file)
                        dest_path = os.path.join(lost_photos_dir, f"{turma_dir}_{photo_file}")
                        try:
                            shutil.move(source_path, dest_path)
                            orphaned_photos_count += 1
                        except Exception as e:
                            print(f"Erro ao mover foto órfã {source_path}: {e}")
    
    # Processar thumbnails
    for turma_dir in os.listdir(THUMBS_DIR) if os.path.exists(THUMBS_DIR) else []:
        if turma_dir == 'lost':  # Pular a pasta lost
            continue
            
        turma_path = os.path.join(THUMBS_DIR, turma_dir)
        if os.path.isdir(turma_path):
            for thumb_file in os.listdir(turma_path):
                if thumb_file.endswith('.jpg'):
                    processo = thumb_file[:-4]  # Remove .jpg
                    aluno = Aluno.query.filter_by(processo=processo).first()
                    if not aluno:
                        # Thumbnail órfã - mover para pasta lost
                        source_path = os.path.join(turma_path, thumb_file)
                        dest_path = os.path.join(lost_thumbs_dir, f"{turma_dir}_{thumb_file}")
                        try:
                            shutil.move(source_path, dest_path)
                            orphaned_thumbs_count += 1
                        except Exception as e:
                            print(f"Erro ao mover thumbnail órfã {source_path}: {e}")
    
    db.session.commit()
    
    # Construir mensagem personalizada baseada na opção escolhida e fotos órfãs encontradas
    messages = []
    if mark_as_taken:
        messages.append(f'Rescan concluído: {updated_count} alunos marcados com foto existente e tirada.')
    else:
        messages.append(f'Rescan concluído: {updated_count} alunos marcados com foto existente mas não tirada.')
    
    if orphaned_photos_count > 0 or orphaned_thumbs_count > 0:
        messages.append(f'Fotos órfãs movidas para pasta "lost": {orphaned_photos_count} fotos originais e {orphaned_thumbs_count} thumbnails.')
    
    flash(' '.join(messages), 'success')
    
    return redirect(url_for('settings'))
# End def settings_rescan_photos


@app.route('/settings/csv/', methods=['POST'])
@required_login
@required_role('admin')
def settings_csv():
    """Upload CSV com substituição ou merge dos dados baseado no campo action"""
    action = request.form.get('action', 'replace')
    
    if 'file' not in request.files:
        flash('Nenhum ficheiro foi selecionado!', 'error')
        return redirect(url_for('settings'))

    file = request.files['file']
    if file.filename == '':
        flash('Nenhum ficheiro foi selecionado!', 'error')
        return redirect(url_for('settings'))

    # Ação específica para atualizar professores
    if action == 'update_professor':
        try:
            # Ler o conteúdo do ficheiro diretamente da memória
            csv_content = file.read().decode('utf-8')
            
            # Usar StringIO para simular um ficheiro
            from io import StringIO
            csv_file = StringIO(csv_content)
            
            # Ler CSV e processar dados
            reader = csv.DictReader(csv_file)
            turmas_atualizadas = 0
            turmas_nao_encontradas = []
            
            for row in reader:
                turma_nome = row.get('turma', '').strip()
                professor_nome = row.get('professor', '').strip()
                
                if not turma_nome:
                    continue  # Pular linhas sem nome de turma
                
                # Buscar turma na base de dados
                turma = Turma.query.filter_by(nome=turma_nome).first()
                
                if turma:
                    # Atualizar nome do professor
                    turma.nome_professor = professor_nome
                    
                    # Atualizar email do professor se disponível no CSV
                    # Verificar diferentes possíveis nomes de coluna para email
                    email_professor = ''
                    if 'email_professor' in row and row['email_professor']:
                        email_professor = row['email_professor'].strip()
                    elif 'email' in row and row['email']:
                        email_professor = row['email'].strip()
                    
                    if email_professor:
                        turma.email_professor = email_professor
                    
                    turmas_atualizadas += 1
                else:
                    # Turma não encontrada - adicionar à lista para flash
                    turmas_nao_encontradas.append(turma_nome)
            
            # Commit das alterações
            db.session.commit()
            
            # Construir mensagens de feedback
            messages = []
            if turmas_atualizadas > 0:
                messages.append(f'{turmas_atualizadas} turma(s) atualizada(s) com sucesso.')
            
            if turmas_nao_encontradas:
                turmas_ignoradas = ', '.join(turmas_nao_encontradas[:5])  # Limitar a 5 nomes
                if len(turmas_nao_encontradas) > 5:
                    turmas_ignoradas += f' (e mais {len(turmas_nao_encontradas) - 5})'
                messages.append(f'Turmas ignoradas (não encontradas): {turmas_ignoradas}.')
            
            flash(' '.join(messages), 'success' if turmas_atualizadas > 0 else 'info')
            
        except Exception as e:
            db.session.rollback()
            print(f"Erro ao atualizar professores: {e}")
            flash(f'Erro ao processar ficheiro CSV para atualização de professores: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    # Ações para substituir/merge de dados completos
    replace_all = (action == 'replace')

    if file:
        try:
            # Ler o conteúdo do ficheiro diretamente da memória
            csv_content = file.read().decode('utf-8')
            
            # Usar StringIO para simular um ficheiro
            from io import StringIO
            csv_file = StringIO(csv_content)
            
            # Mapear fotos existentes antes das alterações
            existing_photos = {}
            for turma_dir in os.listdir(PHOTOS_DIR) if os.path.exists(PHOTOS_DIR) else []:
                turma_path = os.path.join(PHOTOS_DIR, turma_dir)
                if os.path.isdir(turma_path):
                    for photo_file in os.listdir(turma_path):
                        if photo_file.endswith('.jpg'):
                            processo = photo_file[:-4]  # Remove .jpg
                            existing_photos[processo] = turma_dir
            
            # Mapear thumbnails existentes
            existing_thumbs = {}
            for turma_dir in os.listdir(THUMBS_DIR) if os.path.exists(THUMBS_DIR) else []:
                turma_path = os.path.join(THUMBS_DIR, turma_dir)
                if os.path.isdir(turma_path):
                    for thumb_file in os.listdir(turma_path):
                        if thumb_file.endswith('.jpg'):
                            processo = thumb_file[:-4]  # Remove .jpg
                            existing_thumbs[processo] = turma_dir
            
            if replace_all:
                # Limpar dados existentes da base de dados (mas manter fotos)
                db.session.query(Aluno).delete()
                db.session.query(Turma).delete()
                turmas_dict = {}
            else:
                # Para merge, carregar turmas existentes
                turmas_dict = {turma.nome: turma for turma in Turma.query.order_by(Turma.nome).all()}
            
            # Ler CSV e importar dados
            reader = csv.DictReader(csv_file)
            photos_moved = 0
            photos_kept = 0
            alunos_processados = 0
            processos_csv = set()  # Controlar duplicatas no CSV
            
            for row in reader:
                turma_nome = row['turma']
                processo_original = row['processo']
                
                # Validar e sanitizar processo
                processo, erro_processo = Aluno.validate_and_sanitize_processo(processo_original)
                if erro_processo:
                    flash(f'Erro no CSV - linha com turma "{turma_nome}": {erro_processo}', 'error')
                    return redirect(url_for('settings'))
                
                # Verificar duplicata no próprio CSV
                if processo in processos_csv:
                    flash(f'Erro no CSV: Processo {processo} aparece duplicado no ficheiro!', 'error')
                    return redirect(url_for('settings'))
                processos_csv.add(processo)
                
                # Criar turma se não existir
                if turma_nome not in turmas_dict:
                    turma = Turma(nome=turma_nome)
                    db.session.add(turma)
                    db.session.flush()  # Para obter o ID
                    turmas_dict[turma_nome] = turma
                else:
                    turma = turmas_dict[turma_nome]
                
                # Verificar se aluno já existe globalmente (para merge)
                if not replace_all:
                    aluno_existente = Aluno.query.filter_by(processo=processo).first()
                    
                    if aluno_existente:
                        # Se o aluno existe mas está numa turma diferente, é um erro
                        if aluno_existente.turma.nome != turma_nome:
                            flash(f'Erro no CSV: Aluno com processo {processo} já existe na turma "{aluno_existente.turma.nome}" mas CSV indica turma "{turma_nome}"!', 'error')
                            return redirect(url_for('settings'))
                        
                        # Atualizar dados do aluno existente na mesma turma
                        aluno_existente.nome = row['nome']
                        # Atualizar número se disponível
                        if 'numero' in row and row['numero']:
                            try:
                                aluno_existente.numero = int(row['numero'])
                            except ValueError:
                                pass  # Ignorar se não for um número válido
                        # Atualizar email se disponível no CSV
                        if 'email' in row:
                            aluno_existente.email = row['email'].strip()
                        alunos_processados += 1
                        continue
                
                # Preparar número
                numero_val = None
                if 'numero' in row and row['numero']:
                    try:
                        numero_val = int(row['numero'])
                    except ValueError:
                        pass  # Ignorar se não for um número válido
                
                # Verificar se existe foto para este processo
                foto_tirada = False
                if processo in existing_photos:
                    foto_tirada = True
                    old_turma = existing_photos[processo]
                    
                    # Se a foto está numa turma diferente, mover
                    if old_turma != turma_nome:
                        try:
                            # Mover foto
                            old_photo_path = os.path.join(PHOTOS_DIR, old_turma, f'{processo}.jpg')
                            new_photo_path = os.path.join(PHOTOS_DIR, turma_nome, f'{processo}.jpg')
                            
                            # Criar diretório de destino se não existir
                            safe_makedirs(os.path.dirname(new_photo_path))
                            
                            if os.path.exists(old_photo_path):
                                import shutil
                                shutil.move(old_photo_path, new_photo_path)
                                photos_moved += 1
                            
                            # Mover thumbnail se existir
                            if processo in existing_thumbs:
                                old_thumb_path = os.path.join(THUMBS_DIR, old_turma, f'{processo}.jpg')
                                new_thumb_path = os.path.join(THUMBS_DIR, turma_nome, f'{processo}.jpg')
                                
                                safe_makedirs(os.path.dirname(new_thumb_path))
                                
                                if os.path.exists(old_thumb_path):
                                    shutil.move(old_thumb_path, new_thumb_path)
                        
                        except Exception as e:
                            print(f"Erro ao mover foto do processo {processo}: {e}")
                    else:
                        photos_kept += 1
                
                # Preparar email se disponível no CSV
                email_val = ''
                if 'email' in row and row['email']:
                    email_val = row['email'].strip()
                
                # Criar novo aluno
                aluno = Aluno(
                    processo=processo,
                    nome=row['nome'],
                    numero=numero_val,
                    email=email_val,
                    turma_id=turma.id,
                    foto_tirada=foto_tirada
                )
                db.session.add(aluno)
                alunos_processados += 1
            
            db.session.commit()
            
            # Construir mensagem de sucesso detalhada
            messages = []
            if replace_all:
                messages.append(f'Dados substituídos com sucesso! {alunos_processados} alunos processados.')
            else:
                messages.append(f'Dados atualizados com sucesso! {alunos_processados} alunos processados.')
            
            if photos_moved > 0:
                messages.append(f'{photos_moved} foto(s) movida(s) para nova(s) turma(s).')
            
            if photos_kept > 0:
                messages.append(f'{photos_kept} foto(s) mantida(s) na turma correta.')
            
            # Adicionar informação sobre continuar
            messages.append('Pode agora visualizar as turmas ou carregar mais dados.')
            
            flash(' '.join(messages), 'success')
            return redirect(url_for('settings'))
            
        except Exception as e:
            db.session.rollback()
            print(f"Erro ao processar CSV: {e}")
            flash(f'Erro ao processar ficheiro CSV: {str(e)}', 'error')
            return redirect(url_for('settings'))
# End def settings_csv


@app.route('/settings/nuke/', methods=['POST'])
@required_login
@required_role('admin')
def settings_nuke():
    """Limpeza completa de todos os dados e fotos - requer permissão de administrador"""
    try:
        # Remover todos os dados da base de dados
        db.session.query(Aluno).delete()
        db.session.query(Turma).delete()
        db.session.commit()
        
        # Remover todos os diretórios de fotos e thumbnails
        import shutil
        
        if os.path.exists(PHOTOS_DIR):
            shutil.rmtree(PHOTOS_DIR)
            safe_makedirs(PHOTOS_DIR)
        
        if os.path.exists(THUMBS_DIR):
            shutil.rmtree(THUMBS_DIR)
            safe_makedirs(THUMBS_DIR)
        
        flash('Limpeza completa realizada com sucesso! Todos os dados e fotos foram removidos.', 'success')
        
    except Exception as e:
        db.session.rollback()
        print(f"Erro na limpeza completa: {e}")
        flash(f'Erro ao realizar limpeza completa: {str(e)}', 'error')
    
    return redirect(url_for('settings'))
# End def settings_nuke


@app.route('/settings/backup/', methods=['POST'])
@required_login
@required_role('admin')
def settings_backup():
    """Criar backup completo da base de dados em formato ZIP com ficheiros CSV"""
    try:
        # Criar um ZIP em memória
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # 1. Backup dos utilizadores (users.csv)
            users_csv = StringIO()
            users_writer = csv.writer(users_csv)
            users_writer.writerow(['username', 'password', 'name', 'role', 'is_verified'])
            
            users = User.query.order_by(User.id).all()
            for user in users:
                users_writer.writerow([
                    user.email,  # username
                    user.password_hash,  # password hash
                    user.name,
                    user.role,
                    user.is_verified
                ])
            
            zip_file.writestr('users.csv', users_csv.getvalue())
            
            # 2. Backup dos alunos (alunos.csv)
            alunos_csv = StringIO()
            alunos_writer = csv.writer(alunos_csv)
            alunos_writer.writerow(['processo', 'turma', 'numero', 'nome', 'email', 'autorizacao', 'notes'])
            
            alunos = Aluno.query.order_by(Aluno.id).all()
            for aluno in alunos:
                alunos_writer.writerow([
                    aluno.processo,
                    aluno.turma.nome,  # Nome da turma em vez do ID
                    aluno.numero if aluno.numero else '',
                    aluno.nome,
                    aluno.email,
                    aluno.autorizacao,
                    aluno.notes if aluno.notes else '',
                ])
            
            zip_file.writestr('alunos.csv', alunos_csv.getvalue())
            
            # 3. Backup dos professores (turmas_professores.csv)
            professores_csv = StringIO()
            professores_writer = csv.writer(professores_csv)
            professores_writer.writerow(['turma', 'professor', 'email'])
            
            turmas = Turma.query.order_by(Turma.id).all()
            for turma in turmas:
                professores_writer.writerow([
                    turma.nome,
                    turma.nome_professor,
                    turma.email_professor
                ])
            
            zip_file.writestr('turmas_professores.csv', professores_csv.getvalue())
        
        zip_buffer.seek(0)
        
        # Criar nome do ficheiro com data/hora
        from datetime import datetime
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'backup_database_{timestamp}.zip'
        
        # Retornar o ficheiro ZIP como download
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/zip'
        )
        
    except Exception as e:
        print(f"Erro ao criar backup: {e}")
        flash(f'Erro ao criar backup da base de dados: {str(e)}', 'error')
        return redirect(url_for('settings'))
# End def settings_backup


# Rotas para gestão de utilizadores integradas em /settings
@app.route('/settings/users/', methods=['POST'])
@app.route('/settings/users/<int:user_id>/', methods=['POST'])
@required_login
@required_role('admin')
def user_management(user_id=None):
    """Gestão completa de utilizadores baseada em actions"""
    action = request.form.get('action', '').strip()
    
    if action == 'crud_user_add_new':
        """Criar novos utilizadores"""
        email = request.form.get('email', '').strip()
        name = request.form.get('name', '').strip()
        password = request.form.get('password', '')
        role = request.form.get('role', 'none').strip()
        
        # Validar campos obrigatórios
        if not email or not name or not password:
            flash('Email, nome e password são obrigatórios!', 'error')
            return redirect(url_for('settings'))
        
        # Validar formato do email
        if not AddUserSecurityCheck.validate_email_format(email):
            flash('Formato de email inválido.', 'error')
            return redirect(url_for('settings'))
        
        # Validar força da password
        is_strong, password_error = AddUserSecurityCheck.validate_password_strength(password)
        if not is_strong:
            flash(password_error, 'error')
            return redirect(url_for('settings'))
        
        # Validar role
        valid_roles = ['none', 'viewer', 'editor', 'admin']
        if role not in valid_roles:
            flash('Role inválido.', 'error')
            return redirect(url_for('settings'))
        
        # Verificar se o email já existe
        existing_user = User.query.filter_by(_email=email).first()
        if existing_user:
            flash('Já existe um utilizador com este email.', 'error')
            return redirect(url_for('settings'))
        
        try:
            # Criar novo utilizador
            new_user = User()
            new_user._email = email  # Set directly to avoid validation during construction
            new_user.name = name
            new_user.password = password  # Uses the setter that hashes the password
            new_user.role = role
            new_user.is_verified = True  # Admin-created users are pre-verified
            
            db.session.add(new_user)
            db.session.commit()
            
            flash(f'Utilizador "{name}" criado com sucesso com role "{role}"!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao criar utilizador: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    elif action == 'crud_user_edit':
        """Editar dados de um utilizador"""
        if user_id is None:
            flash('ID do utilizador não fornecido para edição.', 'error')
            return redirect(url_for('settings'))
            
        user = User.query.get_or_404(user_id)
        current_admin = get_current_user()
        
        # Não permitir que um admin se remova a si próprio do papel de admin
        if user.id == current_admin.id and request.form.get('role') != 'admin':
            flash('Não pode remover o seu próprio papel de administrador.', 'error')
            return redirect(url_for('settings'))
        
        try:
            # Atualizar dados do utilizador
            user.name = request.form.get('name', '').strip()
            user.email = request.form.get('email', '').strip()
            user.role = request.form.get('role', 'none')
            
            # Verificar se deve notificar o utilizador
            notify_user = request.form.get('notify_user') == 'on' or request.form.get('notify_user') == 'true'

            # Validações básicas
            if not user.name or not user.email:
                flash('Nome e email são obrigatórios.', 'error')
                return redirect(url_for('settings'))
            
            # Verificar se email já existe (exceto para o próprio utilizador)
            existing_email = User.query.filter(User._email == user.email, User.id != user_id).first()
            if existing_email:
                flash('Email já existe.', 'error')
                return redirect(url_for('settings'))
            
            db.session.commit()
            
            # Enviar email de notificação se solicitado
            job_id = None
            
            if notify_user:
                # Enviar email de notificação de forma assíncrona
                from tasks import send_account_updated_email
                
                # Preparar configurações do app para a tarefa
                app_config = {
                    'MAIL_SERVER': app.config['MAIL_SERVER'],
                    'MAIL_PORT': app.config['MAIL_PORT'],
                    'MAIL_USE_TLS': app.config['MAIL_USE_TLS'],
                    'MAIL_USE_SSL': app.config['MAIL_USE_SSL'],
                    'MAIL_USERNAME': app.config['MAIL_USERNAME'],
                    'MAIL_PASSWORD': app.config['MAIL_PASSWORD'],
                    'MAIL_DEFAULT_SENDER': app.config['MAIL_DEFAULT_SENDER'],
                    'TEMPLATES_AUTO_RELOAD': True
                }
                
                # Enqueue da tarefa
                job = email_queue.enqueue(
                    send_account_updated_email,
                    app_config,
                    user.email,
                    user.name,
                    job_timeout=300
                )
                job_id = job.get_id()
            
            # Mensagem de sucesso baseada no envio do email
            if notify_user:
                flash(f'Utilizador {user.name} atualizado com sucesso. Email de notificação está a ser enviado em background.', 'success')
            else:
                flash(f'Utilizador {user.name} atualizado com sucesso.', 'success')
            
            return redirect(url_for('settings'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao atualizar utilizador: {str(e)}', 'error')
            return redirect(url_for('settings'))
    
    elif action == 'crud_user_delete':
        """Eliminar um utilizador"""
        if user_id is None:
            flash('ID do utilizador não fornecido para eliminação.', 'error')
            return redirect(url_for('settings'))
            
        user = User.query.get_or_404(user_id)
        current_admin = get_current_user()
        
        # Não permitir que um admin se elimine a si próprio
        if user.id == current_admin.id:
            flash('Não pode eliminar a sua própria conta.', 'error')
            return redirect(url_for('settings'))
        
        try:
            name = user.name
            # Eliminar também da tabela PreUser se existir
            preuser = PreUser.query.filter_by(_email=user.email).first()
            if preuser:
                db.session.delete(preuser)
            db.session.delete(user)
            db.session.commit()
            flash(f'Utilizador {name} eliminado com sucesso.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao eliminar utilizador: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    elif action == 'crud_user_reset_password':
        """Resetar password de um utilizador"""
        if user_id is None:
            flash('ID do utilizador não fornecido para reset de password.', 'error')
            return redirect(url_for('settings'))
            
        user = User.query.get_or_404(user_id)
        new_password = request.form.get('new_password', '')
        
        if not new_password:
            flash('Nova password é obrigatória!', 'error')
            return redirect(url_for('settings'))
        
        # Validar força da password
        is_strong, password_error = AddUserSecurityCheck.validate_password_strength(new_password)
        if not is_strong:
            flash(password_error, 'error')
            return redirect(url_for('settings'))
        
        try:
            # Atualizar password
            user.password = new_password  # Uses the setter that hashes the password
            db.session.commit()
            flash(f'Password do utilizador "{user.name}" alterada com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao alterar password: {str(e)}', 'error')
        
        return redirect(url_for('settings'))

    elif action == 'crud_preuser_delete':
        if user_id is None:
            flash('ID do utilizador não fornecido para eliminação.', 'error')
            return redirect(url_for('settings'))

        preuser = PreUser.query.get_or_404(user_id)

        try:
            db.session.delete(preuser)
            db.session.commit()
            flash(f'Pré-utilizador {preuser.email} eliminado com sucesso.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao eliminar pré-utilizador: {str(e)}', 'error')
        return redirect(url_for('settings'))

    # Se a ação não for reconhecida, redirecionar
    flash('Ação não reconhecida!', 'error')
    return redirect(url_for('settings'))
# End def user_management » settings_users_crud


# Rota para gestão de IPs banidos integradas em /settings
@app.route('/settings/banned_ips/<int:banned_ip_id>/', methods=['POST'])
@app.route('/settings/banned_ips/clear_all/', methods=['POST'])
@required_login
@required_role('admin')
def banned_ip_management(banned_ip_id=None):
    """Gestão de IPs banidos baseada em actions"""
    action = request.form.get('action', '').strip()
    
    if action == 'crud_banned_ip_delete':
        """Desbanir um IP (eliminar da tabela BannedIPs)"""
        if banned_ip_id is None:
            flash('ID do IP banido não fornecido para desbloqueio.', 'error')
            return redirect(url_for('settings'))
        
        # Buscar IP banido
        banned_ip = db.session.get(BannedIPs, banned_ip_id)
        if not banned_ip:
            flash('IP banido não encontrado!', 'error')
            return redirect(url_for('settings'))
        
        try:
            ip_address = banned_ip.remote_addr
            db.session.delete(banned_ip)
            db.session.commit()
            flash(f'IP {ip_address} desbloqueado com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao desbloquear IP: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    elif action == 'crud_banned_ip_clear_all':
        """Desbanir todos os IPs (limpar tabela BannedIPs)"""
        try:
            # Contar quantos IPs serão desbloqueados
            count = BannedIPs.query.count()
            
            if count == 0:
                flash('Não há IPs banidos para desbloquear.', 'info')
                return redirect(url_for('settings'))
            
            # Eliminar todos os IPs banidos
            BannedIPs.query.delete()
            db.session.commit()
            flash(f'{count} IP{"s" if count != 1 else ""} desbloqueado{"s" if count != 1 else ""} com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao desbloquear todos os IPs: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    # Se a ação não for reconhecida, redirecionar
    flash('Ação não reconhecida para gestão de IPs banidos!', 'error')
    return redirect(url_for('settings'))
# End def banned_ip_management


# Rota para gestão de LoginLog integradas em /settings
@app.route('/settings/login_logs/<int:log_id>/', methods=['POST'])
@app.route('/settings/login_logs/clear_all/', methods=['POST'])
@required_login
@required_role('admin')
def login_log_management(log_id=None):
    """Gestão de logs de login baseada em actions"""
    action = request.form.get('action', '').strip()
    
    if action == 'crud_login_delete':
        """Eliminar um log de login específico"""
        if log_id is None:
            flash('ID do log de login não fornecido para eliminação.', 'error')
            return redirect(url_for('settings'))
        
        # Buscar log de login
        log = db.session.get(LoginLog, log_id)
        if not log:
            flash('Log de login não encontrado!', 'error')
            return redirect(url_for('settings'))
        
        try:
            db.session.delete(log)
            db.session.commit()
            flash(f'Log de login #{log_id} eliminado com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao eliminar log de login: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    elif action == 'crud_login_clear_all':
        """Eliminar todos os logs de login"""
        try:
            # Contar quantos logs serão eliminados
            count = LoginLog.query.count()
            
            if count == 0:
                flash('Não há logs de login para eliminar.', 'info')
                return redirect(url_for('settings'))
            
            # Eliminar todos os logs de login
            LoginLog.query.delete()
            db.session.commit()
            flash(f'{count} log{"s" if count != 1 else ""} de login eliminado{"s" if count != 1 else ""} com sucesso!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao eliminar todos os logs de login: {str(e)}', 'error')
        
        return redirect(url_for('settings'))
    
    # Se a ação não for reconhecida, redirecionar
    flash('Ação não reconhecida para gestão de logs de login!', 'error')
    return redirect(url_for('settings'))
# End def login_log_management


# Obter informações do servidor Redis
@app.route('/settings/redis/server.json', methods=['GET'])
@required_login
@required_role('admin')
def get_redis_status():
    """Rota unificada para Redis - server info + sessions"""
    show_all = (request.args.get('data', '').lower() == 'all')
    
    try:
        client = app.config['SESSION_REDIS']
        
        # Informações do servidor
        t0 = time.time()
        pong = client.ping()
        latency_ms = round((time.time() - t0) * 1000, 2)
        
        info = client.info()
        
        # Contar sessões
        prefix = app.config.get('SESSION_KEY_PREFIX', 'session:')
        pattern = f"{prefix}*"
        sessions_count = 0
        session_keys = []
        
        for key in client.scan_iter(match=pattern, count=1000):
            sessions_count += 1
            session_keys.append(key.decode() if isinstance(key, bytes) else str(key))
        
        # Total de chaves
        total_keys = sum(
            dbinfo.get('keys', 0)
            for name, dbinfo in info.items()
            if isinstance(dbinfo, dict) and name.startswith('db')
        )
        
        payload = {
            "status": "online" if pong else "offline",
            "latency_ms": latency_ms,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "sessions_count": sessions_count,
            "connected_clients": info.get("connected_clients"),
            "used_memory_human": info.get("used_memory_human"),
            "redis_version": info.get("redis_version"),
            "uptime_days": info.get("uptime_in_days"),
            "total_keys": total_keys,
        }
        
        # Debug info se solicitado
        if show_all:
            payload["info"] = info
            payload["session_keys_sample"] = session_keys[:5]  # Mostrar apenas 5 como exemplo
            
            # Testar deserialização de uma sessão se existir
            if session_keys:
                test_key = session_keys[0]
                raw_data = client.get(test_key)
                try:
                    import pickle
                    deserialized = pickle.loads(raw_data)
                    payload["session_deserialization_test"] = {
                        "success": True,
                        "fields": list(deserialized.keys()) if isinstance(deserialized, dict) else [],
                        "sample_data": str(deserialized)[:200]
                    }
                except Exception as e:
                    payload["session_deserialization_test"] = {
                        "success": False,
                        "error": str(e),
                        "raw_type": str(type(raw_data)),
                        "raw_sample": str(raw_data)[:100] if raw_data else "None"
                    }
        
        return jsonify(payload), 200
        
    except redis.exceptions.RedisError as e:
        return jsonify({"status": "offline", "error": str(e)}), 503
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500
#End def get_redis_status


@app.route('/settings/redis/sessions.json', methods=['GET'])
@required_login
@required_role('admin')
def list_redis_sessions():
    show_all = (request.args.get('data', '').lower() == 'all')
    
    client = app.config['SESSION_REDIS']
    prefix = app.config.get('SESSION_KEY_PREFIX', 'session:')
    pattern = f"{prefix}*"

    def load_session_bytes(b):
        if b is None:
            return {}
        
        # Flask-Session pode usar diferentes formatos dependendo da configuração
        try:
            # Tentar MessagePack primeiro (usado pelo Flask-Session por vezes)
            return msgpack.unpackb(b, raw=False, strict_map_key=False)
        except (ImportError, Exception):
            pass
        
        try:
            # Tentar pickle padrão
            return pickle.loads(b)
        except Exception as e1:
            try:
                # Tentar pickle com protocolo específico
                return pickle.loads(b, encoding='latin1')
            except Exception:
                pass
            
            try:
                # Tentar JSON
                if isinstance(b, bytes):
                    return json.loads(b.decode('utf-8', errors='ignore'))
                else:
                    return json.loads(str(b))
            except Exception as e2:
                # Se tudo falhar, tentar extrair dados manualmente do raw
                try:
                    raw_str = str(b)
                    # Procurar por padrões conhecidos nos dados brutos
                    extracted_data = {}
                    
                    # Procurar user_id no raw data
                    user_id_match = re.search(r'user_id.*?(\d+)', raw_str)
                    if user_id_match:
                        extracted_data['user_id'] = int(user_id_match.group(1))
                    
                    # Procurar csrf_token
                    csrf_match = re.search(r'csrf_token.*?([a-f0-9]{40})', raw_str)
                    if csrf_match:
                        extracted_data['csrf_token'] = csrf_match.group(1)
                    
                    # Procurar user_email
                    email_match = re.search(r'user_email.*?([a-zA-Z0-9@._-]+)', raw_str)
                    if email_match:
                        extracted_data['user_email'] = email_match.group(1)
                    
                    if extracted_data:
                        return extracted_data
                    
                except Exception:
                    pass
                
                # Debug: retornar info sobre o erro
                return {
                    "_raw_sample": repr(b)[:200] if b else "None",
                    "_pickle_error": str(e1),
                    "_json_error": str(e2),
                    "_type": str(type(b)),
                    "_length": len(b) if b else 0
                }

    sessions = []
    user_ids = set()
    debug_info = {"total_keys_found": 0, "deserialization_errors": 0}

    for key in client.scan_iter(match=pattern, count=1000):
        debug_info["total_keys_found"] += 1
        key_str = key.decode() if isinstance(key, bytes) else str(key)
        raw = client.get(key)
        data = load_session_bytes(raw)

        # Debug: verificar se conseguimos deserializar
        if "_pickle_error" in data or "_json_error" in data:
            debug_info["deserialization_errors"] += 1

        # Procurar user_id em vários campos possíveis
        user_id = None
        for field in ['user_id', '_user_id', 'current_user_id']:
            if field in data and data[field] is not None:
                user_id = str(data[field])
                break
        
        # Também tentar extrair de objetos user
        if user_id is None and 'user' in data and isinstance(data['user'], dict):
            if 'id' in data['user']:
                user_id = str(data['user']['id'])

        if user_id is not None:
            user_ids.add(user_id)

        ttl = client.ttl(key)
        expires_at = None
        if ttl and ttl > 0:
            expires_at = (datetime.now(timezone.utc) + timedelta(seconds=ttl)).isoformat()

        session_info = {
            "key": key_str.rsplit(':', 1)[-1],  # só o SID
            "user_id": user_id,
            "ttl_seconds": ttl,
            "expires_at": expires_at,
            "last_seen": data.get('last_seen') or '1900-01-01',  # Default para ordenação
            "sliding_expiry": data.get('sliding_expiry', False),  # Mostrar se tem renovação automática
            # Debug: mostrar campos disponíveis na sessão
            "session_fields": list(data.keys()) if isinstance(data, dict) else []
        }
        
        # Em modo debug, incluir dados raw (apenas para desenvolvimento)
        if app.config.get('DEBUG', False):
            session_info["raw_data"] = data

        sessions.append(session_info)

    # Ordenar sessões por last_seen (mais recentes primeiro) - agora com tratamento de None
    sessions.sort(key=lambda x: x.get('last_seen') or '1900-01-01', reverse=True)
    
    # Limitar resultados se não for show_all
    if not show_all:
        sessions = sessions[:10]

    # Enriquecer com informações do utilizador
    users_info = {}
    if user_ids:
        rows = User.query.filter(User.id.in_(list(user_ids))).all()
        users_info = {str(u.id): {"id": u.id, "name": u.name, "email": u.email} for u in rows}

    # Criar lista simplificada apenas com user_id e email (sem duplicatas por utilizador)
    active_users = []
    seen_users = set()
    
    for s in sessions:
        uid = s.get("user_id")
        if uid and uid in users_info and uid not in seen_users:
            user_data = users_info[uid]
            user_sessions = [sess for sess in sessions if sess.get("user_id") == uid]
            
            # Encontrar o last_seen mais recente para este utilizador
            most_recent_last_seen = None
            most_recent_expires_at = None
            for sess in user_sessions:
                last_seen = sess.get("last_seen")
                expires_at = sess.get("expires_at")
                
                # Debug: guardar expires_at mais recente também
                if expires_at and expires_at != 'null':
                    if most_recent_expires_at is None or expires_at > most_recent_expires_at:
                        most_recent_expires_at = expires_at
                
                if last_seen and last_seen != '1900-01-01':
                    if most_recent_last_seen is None or last_seen > most_recent_last_seen:
                        most_recent_last_seen = last_seen
            
            active_users.append({
                "user_id": uid,
                "email": user_data["email"],
                "name": user_data["name"],
                "sessions_count": len(user_sessions),
                "last_seen": most_recent_last_seen or 'Never',
                "expires_at": most_recent_expires_at or 'Never'
            })
            seen_users.add(uid)

    # Ordenar active_users por last_seen (mais recente primeiro)
    active_users.sort(key=lambda x: x.get('last_seen', '1900-01-01'), reverse=True)

    # Criar resposta estruturada com nomes simplificados
    response = {
        "sessions_total": len(sessions),
        "authenticated_users_total": len(active_users),
        "showing_all": show_all,
        "limit_applied": not show_all
    }
    
    # Adicionar utilizadores com nomenclatura simplificada
    for index, user in enumerate(active_users):
        response[f"User {index} last_seen"] = user["last_seen"]
        response[f"User {index} expires_at"] = user["expires_at"]
        response[f"User {index} email"] = user["email"]
        response[f"User {index} name"] = user["name"]
        response[f"User {index} sessions_count"] = user["sessions_count"]
        response[f"User {index} user_id"] = user["user_id"]

    return jsonify(response), 200
# End def list_redis_sessions


# não usada de momento
@app.route('/settings/redis/cleanup', methods=['POST'])
@required_login
@required_role('admin')
def cleanup_redis_sessions():
    """Limpeza manual de sessões Redis expiradas ou inválidas"""
    try:
        client = app.config['SESSION_REDIS']
        prefix = app.config.get('SESSION_KEY_PREFIX', 'session:')
        pattern = f"{prefix}*"
        
        cleaned_sessions = 0
        total_sessions = 0
        
        for key in client.scan_iter(match=pattern, count=1000):
            total_sessions += 1
            
            # Verificar TTL
            ttl = client.ttl(key)
            if ttl == -1:  # Sessão sem TTL (problemática)
                client.delete(key)
                cleaned_sessions += 1
                continue
            
            if ttl == -2:  # Chave já expirou
                cleaned_sessions += 1
                continue
            
            # Verificar se sessão tem dados válidos
            try:
                raw_data = client.get(key)
                if not raw_data:
                    client.delete(key)
                    cleaned_sessions += 1
            except Exception:
                # Se não conseguir ler, remove
                client.delete(key)
                cleaned_sessions += 1
        
        return jsonify({
            "success": True,
            "message": f"Limpeza concluída: {cleaned_sessions} sessões removidas de {total_sessions} analisadas",
            "cleaned_sessions": cleaned_sessions,
            "total_sessions": total_sessions
        }), 200
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Erro na limpeza: {str(e)}"
        }), 500
# End def cleanup_redis_sessions


@app.route('/api/job_status/<job_id>')
@required_login
def job_status(job_id):
    """
    Endpoint para verificar o status de uma tarefa RQ
    """
    try:
        job = Job.fetch(job_id, connection=redis_conn)
        
        # Obter informações do job
        status = job.get_status()
        meta = job.meta or {}
        
        response = {
            'job_id': job_id,
            'status': status,  # 'queued', 'started', 'finished', 'failed'
            'progress': meta.get('progress', 0),
            'message': meta.get('message', ''),
            'error': meta.get('error', '')
        }
        
        # Se a tarefa está concluída, incluir o resultado
        if status == 'finished':
            response['result'] = job.result
        elif status == 'failed':
            response['error'] = str(job.exc_info) if job.exc_info else 'Erro desconhecido'
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({
            'job_id': job_id,
            'status': 'not_found',
            'error': f'Tarefa não encontrada: {str(e)}'
        }), 404
# End function job_status


@app.route('/api/email_jobs')
@required_login
def email_jobs():
    """
    Endpoint para listar todas as tarefas de email na queue
    """
    try:
        # Obter jobs da queue
        jobs = email_queue.get_jobs()
        
        job_list = []
        for job in jobs:
            meta = job.meta or {}
            job_info = {
                'job_id': job.get_id(),
                'status': job.get_status(),
                'progress': meta.get('progress', 0),
                'message': meta.get('message', ''),
                'error': meta.get('error', ''),
                'created_at': job.created_at.isoformat() if job.created_at else None,
                'started_at': job.started_at.isoformat() if job.started_at else None,
                'ended_at': job.ended_at.isoformat() if job.ended_at else None
            }
            job_list.append(job_info)
        
        return jsonify({
            'queue_length': len(email_queue),
            'jobs': job_list
        })
        
    except Exception as e:
        return jsonify({
            'error': f'Erro ao obter tarefas: {str(e)}'
        }), 500
# End function email_jobs


@app.route('/api/send_notification_email', methods=['POST'])
@csrf.exempt  # API endpoint que recebe JSON
@required_login
@required_role('editor')
def send_notification_email():
    """
    API endpoint para enviar email de notificação ao professor responsável
    """
    try:
        # Verificar se é uma chamada AJAX (JSON) ou submissão de formulário
        is_ajax = request.is_json or request.headers.get('Content-Type') == 'application/json'
        
        if is_ajax:
            data = request.get_json()
        else:
            # Se não for AJAX, obter dados do formulário
            data = {
                'email': request.form.get('email', '').strip(),
                'subject': request.form.get('subject', '').strip(),
                'body': request.form.get('body', '').strip(),
                'turma_nome': request.form.get('turma_nome', '').strip(),
                'turma_nome_seguro': request.form.get('turma_nome_seguro', '').strip()
            }
        
        # Validar dados obrigatórios
        required_fields = ['email', 'subject', 'body', 'turma_nome', 'turma_nome_seguro']
        for field in required_fields:
            if field not in data or not data[field]:
                error_msg = f'Campo obrigatório ausente: {field}'
                if is_ajax:
                    return jsonify({'success': False, 'error': error_msg}), 400
                else:
                    flash(error_msg, 'error')
                    return redirect(request.referrer or url_for('turma', nome_seguro=data.get('turma_nome_seguro', '')))
        
        email = data['email'].strip()
        subject = data['subject'].strip()
        body = data['body'].strip()
        turma_nome = data['turma_nome'].strip()
        turma_nome_seguro = data['turma_nome_seguro'].strip()
        
        # Validar formato do email
        if not AddUserSecurityCheck.validate_email_format(email):
            error_msg = 'Formato de email inválido'
            if is_ajax:
                return jsonify({'success': False, 'error': error_msg}), 400
            else:
                flash(error_msg, 'error')
                return redirect(request.referrer or url_for('turma', nome_seguro=turma_nome_seguro))
        
        # Verificar se a turma existe e obter informações
        turma_obj = Turma.query.filter_by(nome_seguro=turma_nome_seguro).first()
        if not turma_obj:
            error_msg = 'Turma não encontrada'
            if is_ajax:
                return jsonify({'success': False, 'error': error_msg}), 404
            else:
                flash(error_msg, 'error')
                return redirect(request.referrer or url_for('turmas'))
        
        # Verificar se o email do professor está configurado
        if not turma_obj.email_professor or turma_obj.email_professor != email:
            error_msg = 'Email do professor não está configurado para esta turma'
            if is_ajax:
                return jsonify({'success': False, 'error': error_msg}), 400
            else:
                flash(error_msg, 'error')
                return redirect(url_for('turma', nome_seguro=turma_nome_seguro))
        
        # Preparar configurações do app para a tarefa
        app_config = {
            'MAIL_SERVER': app.config['MAIL_SERVER'],
            'MAIL_PORT': app.config['MAIL_PORT'],
            'MAIL_USE_TLS': app.config['MAIL_USE_TLS'],
            'MAIL_USE_SSL': app.config['MAIL_USE_SSL'],
            'MAIL_USERNAME': app.config['MAIL_USERNAME'],
            'MAIL_PASSWORD': app.config['MAIL_PASSWORD'],
            'MAIL_DEFAULT_SENDER': app.config['MAIL_DEFAULT_SENDER'],
            'TEMPLATES_AUTO_RELOAD': True
        }
        
        # Importar e executar tarefa de envio
        from tasks import send_notification_teacher_email
        
        # Enqueue da tarefa
        job = email_queue.enqueue(
            send_notification_teacher_email,
            app_config,
            email,
            subject,
            body,
            turma_nome,
            turma_obj.nome_professor or 'Professor',
            job_timeout=300
        )
        
        print(f"Tarefa de notificação criada: {job.get_id()}")
        
        success_msg = 'Email de notificação enviado com sucesso! O processamento será feito em segundo plano.'
        
        if is_ajax:
            return jsonify({
                'success': True,
                'message': success_msg,
                'job_id': job.get_id(),
                'email': email
            })
        else:
            flash(success_msg, 'success')
            return redirect(url_for('turma', nome_seguro=turma_nome_seguro))
        
    except Exception as e:
        print(f"Erro ao processar notificação: {str(e)}")
        error_msg = f'Erro interno do servidor: {str(e)}'
        if request.is_json or request.headers.get('Content-Type') == 'application/json':
            return jsonify({'success': False, 'error': error_msg}), 500
        else:
            flash('Erro ao enviar email. Tente novamente.', 'error')
            return redirect(request.referrer or url_for('turmas'))
# End function send_notification_email




if __name__ == '__main__':
    
    create_directories_with_permissions()
    # Inicializar base de dados quando o script é executado diretamente
    with app.app_context():
        db.create_all()
        print("Base de dados inicializada.")
        
    #app.run(debug=DEBUG, host='0.0.0.0', port=5000)
    app.run(debug=True, host='0.0.0.0', port=5000)
